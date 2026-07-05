"""
Indent PDF service — python-docx + docx2pdf (MS Word COM).

Security layers
---------------
1. Signed QR token  (itsdangerous URLSafeTimedSerializer, 2-day expiry)
2. Tamper fingerprint SHA-256(indent_no|project_code|items) — embedded in
   token AND printed in the PDF body.  Verify endpoint recomputes live and
   compares; mismatch → TAMPERED page.
3. Micro-print footer  4pt repeating text on every page.
4. APPROVED diagonal watermark in Word header (every page, approved docs only).
5. VERIFIED section + diagonal watermark on last page (approved docs only).
6. cantSplit on every item row — no row splits across pages.
7. Repeating header (logo + website + QR) via Word section header.
"""
import io, os, copy, hashlib, tempfile
import qrcode
import pythoncom
from datetime import datetime
from lxml import etree
from PIL import Image, ImageDraw
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from flask import current_app, make_response, send_from_directory

from app.extensions import db
from app.models.indent_master import IndentMaster
from app.models.approval_path import ApprovalHistory
from app.models.user import User
from app.response import res

# ── XML namespace constants ───────────────────────────────────────────────────
W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP   = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A    = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
W10  = "urn:schemas-microsoft-com:office:word"

# ── app constants ─────────────────────────────────────────────────────────────
PDF_TOKEN_SALT  = "indent-pdf-v1"
PDF_EXPIRY_DAYS = 2
TEMPLATE_PATH   = os.path.join(os.path.dirname(__file__), "indent_template.docx")
LOGO_PATH       = os.getenv("COMPANY_LOGO_PATH",
                             os.path.join(os.getcwd(), "asset", "order", "1000018063.jpeg"))


# ═════════════════════════════════════════════════════════════════════════════
# Storage helpers  (swap this block for BunnyCDN / S3)
# ═════════════════════════════════════════════════════════════════════════════

def _storage_root():
    return current_app.config.get(
        "INDENT_PDF_STORAGE_PATH",
        os.path.join(os.getcwd(), "storage", "pdf"),
    )


def _save_pdf(pdf_bytes, indent_no):
    root   = _storage_root()
    folder = os.path.join(root, "indent_pdf", indent_no)
    os.makedirs(folder, exist_ok=True)
    with open(os.path.join(folder, "indent.pdf"), "wb") as f:
        f.write(pdf_bytes)
    base = current_app.config.get("INDENT_PDF_BASE_URL", "/resource/indent/pdf-file")
    return f"{base}/indent_pdf/{indent_no}/indent.pdf"


def serve_indent_pdf(relative_path):
    root     = _storage_root()
    full_dir = os.path.join(root, os.path.dirname(relative_path))
    return send_from_directory(full_dir, os.path.basename(relative_path),
                               mimetype="application/pdf")


# ═════════════════════════════════════════════════════════════════════════════
# Crypto / fingerprint
# ═════════════════════════════════════════════════════════════════════════════

def _serializer():
    secret = current_app.config.get("JWT_SECRET_KEY", "erp-secret")
    return URLSafeTimedSerializer(secret)


def _fingerprint(indent):
    """
    SHA-256 of indent_no + project_code + every item (code:qty sorted).
    32-char uppercase hex.  Embedded in QR token and printed in the PDF.
    Any change to indent data after PDF generation will cause a mismatch.
    """
    items_raw = "|".join(
        f"{(ii.item_code or '').upper()}:{float(ii.qty or 0):.3f}"
        for ii in sorted(indent.indent_items, key=lambda x: x.item_code or "")
    )
    raw = f"{indent.indent_no}|{indent.project_code}|{items_raw}"
    return hashlib.sha256(raw.encode()).hexdigest()[:32].upper()


# ═════════════════════════════════════════════════════════════════════════════
# Approval history
# ═════════════════════════════════════════════════════════════════════════════

def _build_signatures(indent_id):
    rows = (
        db.session.query(ApprovalHistory, User.username)
        .outerjoin(User, User.id == ApprovalHistory.action_by)
        .filter(ApprovalHistory.module_code == "indent",
                ApprovalHistory.record_id   == indent_id)
        .order_by(ApprovalHistory.id.asc())
        .all()
    )
    prepared_by = None
    checked_by  = []
    authorised_by = None
    for h, username in rows:
        name = username or "—"
        at   = h.created_at.strftime("%d-%b-%Y; %H:%M") if h.created_at else ""
        if   h.action == "SUBMIT":        prepared_by   = {"name": name, "at": at}
        elif h.action == "FINAL_APPROVE": authorised_by = {"name": name, "at": at}
        elif h.action == "APPROVE":       checked_by.append({"name": name, "at": at})
    return prepared_by, checked_by, authorised_by


# ═════════════════════════════════════════════════════════════════════════════
# QR code with L-shaped corner brackets
# ═════════════════════════════════════════════════════════════════════════════

def _qr_with_brackets(url):
    """Generate a plain QR PNG then draw L-bracket corners. Returns PNG bytes."""
    from qrcode.image.pil import PilImage

    qr = qrcode.QRCode(box_size=10, border=4,
                       error_correction=qrcode.constants.ERROR_CORRECT_M)
    qr.add_data(url)
    qr.make(fit=True)

    buf = io.BytesIO()
    qr.make_image(image_factory=PilImage,
                  fill_color="black", back_color="white").save(buf, "PNG")
    buf.seek(0)

    img  = Image.open(buf).convert("RGB")
    draw = ImageDraw.Draw(img)
    iw, ih = img.size
    t = max(5, iw // 45)   # bracket arm thickness
    L = iw // 5            # bracket arm length
    B = "black"

    for (x0, y0, x1, y1) in [
        # top-left
        (0,    0,    L,    t   ), (0,    0,    t,    L   ),
        # top-right
        (iw-L, 0,    iw,   t   ), (iw-t, 0,    iw,   L   ),
        # bottom-left
        (0,    ih-t, L,    ih  ), (0,    ih-L, t,    ih  ),
        # bottom-right
        (iw-L, ih-t, iw,   ih  ), (iw-t, ih-L, iw,   ih  ),
    ]:
        draw.rectangle([x0, y0, x1-1, y1-1], fill=B)

    out = io.BytesIO()
    img.save(out, "PNG")
    return out.getvalue()


# ═════════════════════════════════════════════════════════════════════════════
# VML diagonal watermark paragraph element
# ═════════════════════════════════════════════════════════════════════════════

def _watermark_para_el(text, color_hex="D8D8D8"):
    """
    Return a <w:p> lxml element containing a VML diagonal watermark shape.
    Works in both the Word section header (all-pages) and in the body
    (last-page only, anchored to the paragraph it follows).
    color_hex: 6-char hex without '#'.
    """
    xml = (
        f'<w:p xmlns:w="{W}" xmlns:v="{V_NS}"'
        f' xmlns:o="{O_NS}" xmlns:w10="{W10}">'
        f'<w:pPr><w:rPr><w:vanish/></w:rPr></w:pPr>'
        f'<w:r><w:rPr><w:vanish/></w:rPr>'
        f'<w:pict>'
        f'<v:shape id="WaterMark" type="_x0000_t136"'
        f' style="position:absolute;left:0;text-align:center;'
        f'margin-left:0;margin-top:0;'
        f'width:500pt;height:150pt;z-index:-251655168;'
        f'mso-position-horizontal:center;'
        f'mso-position-horizontal-relative:margin;'
        f'mso-position-vertical:center;'
        f'mso-position-vertical-relative:margin;'
        f'rotation:-45"'
        f' filled="t" stroked="f" fillcolor="#{color_hex}">'
        f'<v:fill on="t" focussize="0,0"/>'
        f'<v:textbox style="mso-fit-shape-to-text:t" inset="0,0,0,0">'
        f'<w:txbxContent><w:p><w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'<w:color w:val="{color_hex}"/>'
        f'<w:sz w:val="120"/>'
        f'</w:rPr>'
        f'<w:t>{text}</w:t>'
        f'</w:r></w:p></w:txbxContent>'
        f'</v:textbox>'
        f'<w10:wrap type="none"/>'
        f'<w10:anchorlock/>'
        f'</v:shape>'
        f'</w:pict>'
        f'</w:r></w:p>'
    )
    return etree.fromstring(xml)


# ═════════════════════════════════════════════════════════════════════════════
# Word section header  (repeating on every page)
# ═════════════════════════════════════════════════════════════════════════════

def _build_header(doc, qr_bytes, generated_at):
    """
    Build a 3-column header table [Logo | Website + date | QR] in the
    section header so it repeats on every page.
    Also strips the floating logo / textbox anchors from body P0.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    header = section.header
    header.is_linked_to_previous = False

    # Clear any existing header paragraphs
    hdr_el = header._element
    for child in list(hdr_el):
        hdr_el.remove(child)

    # 1-row × 3-col borderless table
    tbl = header.add_table(rows=1, cols=3, width=Inches(7.2))
    tbl.autofit = False
    tbl.allow_autofit = False

    tbl_pr = tbl._tbl.tblPr
    for b in tbl_pr.findall(f"{{{W}}}tblBorders"):
        tbl_pr.remove(b)
    brd = etree.SubElement(tbl_pr, f"{{{W}}}tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = etree.SubElement(brd, f"{{{W}}}{edge}")
        e.set(f"{{{W}}}val", "none")
        e.set(f"{{{W}}}sz",  "0")

    cells = tbl.rows[0].cells
    cells[0].width = Inches(2.1)   # wider for larger logo
    cells[1].width = Inches(3.0)
    cells[2].width = Inches(2.1)

    # All cells: top-aligned, zero paragraph spacing
    def _cell_top(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(f"{{{W}}}vAlign"):
            tc_pr.remove(old)
        va = etree.SubElement(tc_pr, f"{{{W}}}vAlign")
        va.set(f"{{{W}}}val", "top")
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

    for c in cells:
        _cell_top(c)

    # Left cell — logo only
    if os.path.exists(LOGO_PATH):
        lp = cells[0].paragraphs[0]
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(0)
        lp.add_run().add_picture(LOGO_PATH, width=Inches(2.0))

    # Centre cell — INDENT title (vertically centred, bold, large)
    tc_pr_mid = cells[1]._tc.get_or_add_tcPr()
    for old in tc_pr_mid.findall(f"{{{W}}}vAlign"):
        tc_pr_mid.remove(old)
    va_mid = etree.SubElement(tc_pr_mid, f"{{{W}}}vAlign")
    va_mid.set(f"{{{W}}}val", "center")

    ip = cells[1].paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(0)
    ip.paragraph_format.space_after  = Pt(0)
    ir = ip.add_run("INDENT")
    ir.font.size = Pt(26)
    ir.font.bold = True
    ir.font.name = "Calibri"

    # Right cell layout (cell width = 2.1", starts at 5.10" from page content left)
    # Right values (col3) start at 5.70" → QR left indent = 5.70 - 5.10 = 0.60"
    # QR left edge sits at 5.70": within the ": value" region of the right column.
    _QR_W    = Inches(0.975)
    _QR_LEFT = Inches(1.0)    # right values start exactly at 5.10" = right cell left edge

    # Website text: same left edge as QR, 11pt, left-aligned
    qp = cells[2].paragraphs[0]
    qp.alignment                     = WD_ALIGN_PARAGRAPH.LEFT
    qp.paragraph_format.space_before = Pt(0)
    qp.paragraph_format.space_after  = Pt(5)
    qp.paragraph_format.left_indent  = Inches(-0.05)  # website text only — adjust to push left/right independently of QR
    qp.paragraph_format.right_indent = Pt(0)
    wr = qp.add_run("www.dishaanhitech.com")
    wr.font.size = Pt(12)
    wr.font.bold = True
    wr.font.name = "Bahnschrift Light"

    qr_p = cells[2].add_paragraph()
    qr_p.alignment                     = WD_ALIGN_PARAGRAPH.LEFT
    qr_p.paragraph_format.space_before = Pt(0)
    qr_p.paragraph_format.space_after  = Pt(0)
    qr_p.paragraph_format.left_indent  = _QR_LEFT
    qr_p.add_run().add_picture(io.BytesIO(qr_bytes), width=_QR_W)

    # Word requires the header element to end with a <w:p>
    etree.SubElement(hdr_el, f"{{{W}}}p")

    # ── Body cleanup ──────────────────────────────────────────────────────────
    # Goal: INDENT must be the very first body element.
    #
    # Strategy (do NOT rely on style names):
    #   Pass 1 — strip any <w:r> that carries a <w:drawing> or <w:pict>
    #             (floating logo, website textbox anchors).  Do this first so
    #             that paragraphs which only contained those runs become empty.
    #   Pass 2 — delete EVERY paragraph before INDENT regardless of content.
    #             We own every element before INDENT; none of it belongs in the
    #             body once the header has been built.
    #   Pass 3 — debug log so the caller can verify the result.

    # Pass 1: strip floating runs (drawing + VML pict) from all body paragraphs
    for p_el in list(doc.element.body.findall(f"{{{W}}}p")):
        text_check = "".join(t.text or "" for t in p_el.iter(f"{{{W}}}t")).strip()
        if text_check == "INDENT":
            break
        for r_el in list(p_el.findall(f"{{{W}}}r")):
            if (r_el.find(f"{{{W}}}drawing") is not None
                    or r_el.find(f"{{{W}}}pict")    is not None):
                p_el.remove(r_el)

    # Pass 2: delete every paragraph up to AND including INDENT (now in header)
    for p_el in list(doc.element.body.findall(f"{{{W}}}p")):
        text_check = "".join(t.text or "" for t in p_el.iter(f"{{{W}}}t")).strip()
        p_el.getparent().remove(p_el)
        if text_check == "INDENT":
            break

    # Pass 3: debug log — verify INDENT is now the first body paragraph
    import sys
    print("\n[indent_pdf] Body paragraphs after header cleanup:", file=sys.stderr)
    for i, p_el in enumerate(doc.element.body.findall(f"{{{W}}}p")):
        txt = "".join(t.text or "" for t in p_el.iter(f"{{{W}}}t")).strip()[:60]
        print(f"  P{i} -> {txt!r}", file=sys.stderr)
        if i >= 6:
            print("  ...", file=sys.stderr)
            break


# ═════════════════════════════════════════════════════════════════════════════
# Items table helpers
# ═════════════════════════════════════════════════════════════════════════════

def _set_cant_split(table):
    """Prevent every item row from breaking across pages."""
    for row in table.rows:
        tr_pr = row._tr.find(f"{{{W}}}trPr")
        if tr_pr is None:
            tr_pr = etree.SubElement(row._tr, f"{{{W}}}trPr")
            row._tr.insert(0, tr_pr)
        cs = etree.SubElement(tr_pr, f"{{{W}}}cantSplit")
        cs.set(f"{{{W}}}val", "1")


# ═════════════════════════════════════════════════════════════════════════════
# Micro-print footer
# ═════════════════════════════════════════════════════════════════════════════

def _add_security_footer(doc, indent_no, fp_hash, generated_at=""):
    """Footer: generated timestamp (center) + 4pt micro-print security line."""
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False

    ftr_el = footer._element
    for child in list(ftr_el):
        ftr_el.remove(child)

    # Generated timestamp line
    if generated_at:
        gp = footer.add_paragraph()
        gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gr = gp.add_run(f"Generated: {generated_at}")
        gr.font.size = Pt(8)
        gr.font.name = "Bahnschrift Light"
        gr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 4pt micro-print security line
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    micro = (
        f"DISHAAN HI-TECH (INDIA) PVT LTD | "
        f"INDENT: {indent_no} | HASH: {fp_hash} | AUTHENTIC DOCUMENT | "
    ) * 3
    run = fp.add_run(micro.strip())
    run.font.size = Pt(4)
    run.font.name = "Bahnschrift Light"
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Required closing paragraph
    etree.SubElement(ftr_el, f"{{{W}}}p")


# ═════════════════════════════════════════════════════════════════════════════
# Hash line in body
# ═════════════════════════════════════════════════════════════════════════════

def _add_hash_line(_doc, _fp_hash, _generated_at):
    """Intentionally suppressed — hash/generated info lives only in the footer."""


# ═════════════════════════════════════════════════════════════════════════════
# Verified section (last page, approved docs only)
# ═════════════════════════════════════════════════════════════════════════════

def _add_verified_section(doc, fp_hash):
    """
    Append green ✔ tick + VERIFIED label + hash line to the document body,
    followed by a diagonal 'VERIFIED' watermark anchored to the last paragraph
    so it appears on the last page only.
    """
    GREEN = RGBColor(0x2E, 0x7D, 0x32)

    tick_p = doc.add_paragraph()
    tick_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tick_p.paragraph_format.space_before = Pt(20)
    tr = tick_p.add_run("✔")
    tr.font.size  = Pt(48)
    tr.font.bold  = True
    tr.font.color.rgb = GREEN

    ver_p = doc.add_paragraph("VERIFIED")
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = ver_p.runs[0]
    vr.font.size  = Pt(14)
    vr.font.bold  = True
    vr.font.color.rgb = GREEN

    # Diagonal VERIFIED watermark — anchored after VERIFIED text → last page only
    wm_el = _watermark_para_el("VERIFIED", color_hex="C8E6C9")
    ver_p._p.addnext(wm_el)


# ═════════════════════════════════════════════════════════════════════════════
# Docx fill helpers
# ═════════════════════════════════════════════════════════════════════════════

def _find_para(doc, label):
    for p in doc.paragraphs:
        if p.text.startswith(label):
            return p
    return None


def _t_replace(xml_el, old, new):
    for t in xml_el.iter(f"{{{W}}}t"):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)


def _clear_para(para):
    for r in para.runs:
        r.text = ""


def _set_cell_text(cell, text, bold=False, size_pt=10, font="Bahnschrift Light",
                   space_after_pt=4):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(f"{{{W}}}rFonts")
    if rfonts is None:
        rfonts = etree.SubElement(rpr, f"{{{W}}}rFonts")
    rfonts.set(f"{{{W}}}ascii", font)
    rfonts.set(f"{{{W}}}hAnsi", font)


def _remove_all_borders(table):
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(f"{{{W}}}tblBorders"):
        tblPr.remove(existing)
    borders = etree.SubElement(tblPr, f"{{{W}}}tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = etree.SubElement(borders, f"{{{W}}}{edge}")
        el.set(f"{{{W}}}val",   "none")
        el.set(f"{{{W}}}sz",    "0")
        el.set(f"{{{W}}}space", "0")


def _build_details_table(doc, rows, col_widths_in):
    # 5-column table: col0=left labels, col1=left values,
    #                 col2=SPACER (controls gap between left and right side),
    #                 col3=right labels, col4=right values
    # To control the gap between left values and right labels → adjust col2 width
    # in col_widths_in tuple (index 2). E.g. 0.30" = small gap, 0.60" = large gap.
    table = doc.add_table(rows=len(rows), cols=5)
    table.autofit       = False
    table.allow_autofit = False
    _remove_all_borders(table)
    for r_idx, (ll, lv, rl, rv) in enumerate(rows):
        cells = table.rows[r_idx].cells
        # col0: left label (bold)
        _set_cell_text(cells[0], ll, bold=True,  size_pt=11, space_after_pt=5, font="Bahnschrift Light")
        # col1: left value (regular)
        _set_cell_text(cells[1], lv, bold=False, size_pt=11, space_after_pt=5, font="Bahnschrift Light")
        # col2: spacer — leave empty, width controls gap between left and right side
        _set_cell_text(cells[2], "",  bold=False, size_pt=11, space_after_pt=5, font="Bahnschrift Light")
        # col3: right label (bold)
        _set_cell_text(cells[3], rl, bold=True,  size_pt=11, space_after_pt=5, font="Bahnschrift Light")
        # col4: right value (regular)
        _set_cell_text(cells[4], rv, bold=False, size_pt=11, space_after_pt=5, font="Bahnschrift Light")
    for r in table.rows:
        for ci, w in enumerate(col_widths_in):
            r.cells[ci].width = Inches(w)
    return table


def _duplicate_para_after(ref_para):
    new_p = copy.deepcopy(ref_para._p)
    ref_para._p.addnext(new_p)
    return new_p


# ═════════════════════════════════════════════════════════════════════════════
# Main doc fill  (text fields only — header/QR handled separately)
# ═════════════════════════════════════════════════════════════════════════════

def _fill_doc(doc, indent, prepared_by, checked_by, authorised_by):
    project = getattr(indent, "project", None)

    proj_name     = (project.project_name if project else "") or "—"
    customer      = (project.client_name  if project else "") or "—"
    rsl           = getattr(indent, "site_reg_serial_no", None)
    indent_no_val = indent.indent_no or ""
    if rsl:
        indent_no_val = f"{indent_no_val} /RSL : {rsl}"

    detail_rows = [
        ("Site Code",     f": {indent.project_code or '—'}",
         "Indent No",     f": {indent_no_val}"),
        ("Project Name",  f": {proj_name}",
         "Indent Date",   f": {indent.indent_date or '—'}"),
        ("Customer Name", f": {customer}",
         "Required Date", f": {getattr(indent, 'required_within', '') or '—'}"),
        ("Sale Order No", f": {getattr(indent, 'sale_order_no', '') or '—'}",
         "Indent Category", f": {getattr(indent, 'category_code', '') or '—'}"),
    ]

    # Save original paragraph XML refs BEFORE building the replacement table so
    # _find_para doesn't accidentally pick up table-cell paragraphs afterwards.
    orig_paras = []
    for label in ("Site Code", "Project Name", "Customer Name", "Sale Order No"):
        p = _find_para(doc, label)
        if p is not None:
            orig_paras.append(p._p)

    # Col layout (total 7.20"):
    #   col0=left labels  1.45"  → fits "Customer Name" (longest left label)
    #   col1=left values  2.00"  → ": PC0001", ": ABCD Project" etc.
    #   col2=right labels 1.65"  → fits "Indent Category" / "Required Date"
    #   col3=right values 2.10"  → ": 350015 /RSL : SL123" etc.
    #   col0+col1+col2 = 5.10" → col3 starts exactly where header right cell starts
    #   QR left indent = 0 (flush with right cell)
    # 5-col widths: (left-labels, left-values, GAP-SPACER, right-labels, right-values)
    # Total must always = 7.20"
    # To increase gap between left and right → increase index 2 (spacer), decrease index 1 or 3
    # To shift right labels closer to their colon → decrease index 3
    tbl = _build_details_table(doc, detail_rows, col_widths_in=(1.30, 2.85, 0.25, 1.30, 1.60))
    if orig_paras:
        orig_paras[0].addprevious(tbl._tbl)
    for p_el in orig_paras:
        parent = p_el.getparent()
        if parent is not None:
            parent.remove(p_el)

    # Remove empty paragraphs between items table and audit section to close the gap
    p_placed_check = _find_para(doc, "Indent Placed By")
    if p_placed_check:
        prev = p_placed_check._p.getprevious()
        while prev is not None:
            tag = prev.tag.split("}")[-1] if "}" in prev.tag else prev.tag
            if tag == "p" and not prev.findall(f".//{{{W}}}t"):
                to_remove = prev
                prev = prev.getprevious()
                to_remove.getparent().remove(to_remove)
            else:
                break

    # ── signoff ───────────────────────────────────────────────────────────────
    # All audit rows use the same font, size, and left alignment.
    # Every row is always printed even when the value is missing.
    AUDIT_FONT = "Bahnschrift Light"
    AUDIT_PT   = 11

    def _write_audit_para(para, label, value):
        """Overwrite a template paragraph with consistent label: value formatting."""
        # Clear all existing runs
        for r in list(para.runs):
            r.text = ""
        p_el = para._p
        # Remove every existing <w:r> to start clean
        for r_el in list(p_el.findall(f"{{{W}}}r")):
            p_el.remove(r_el)
        # Left-align, zero spacing
        pPr = p_el.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(p_el, f"{{{W}}}pPr")
            p_el.insert(0, pPr)
        for old in pPr.findall(f"{{{W}}}jc"):
            pPr.remove(old)
        etree.SubElement(pPr, f"{{{W}}}jc").set(f"{{{W}}}val", "left")
        for old in pPr.findall(f"{{{W}}}ind"):
            pPr.remove(old)
        sp = pPr.find(f"{{{W}}}spacing")
        if sp is None:
            sp = etree.SubElement(pPr, f"{{{W}}}spacing")
        sp.set(f"{{{W}}}before", "0")
        sp.set(f"{{{W}}}after",  "40")
        # Tab stop at 1.8" (2592 twips) — colon lands here regardless of label length
        for old in pPr.findall(f"{{{W}}}tabs"):
            pPr.remove(old)
        tabs_el = etree.SubElement(pPr, f"{{{W}}}tabs")
        tab_el  = etree.SubElement(tabs_el, f"{{{W}}}tab")
        tab_el.set(f"{{{W}}}val", "left")
        tab_el.set(f"{{{W}}}pos", "2592")   # 1440 twips/inch × 1.8" = 2592
        # Helper: append one run
        def _run(text, bold=False):
            r = etree.SubElement(p_el, f"{{{W}}}r")
            rPr = etree.SubElement(r, f"{{{W}}}rPr")
            rFonts = etree.SubElement(rPr, f"{{{W}}}rFonts")
            rFonts.set(f"{{{W}}}ascii", AUDIT_FONT)
            rFonts.set(f"{{{W}}}hAnsi", AUDIT_FONT)
            if bold:
                etree.SubElement(rPr, f"{{{W}}}b")
            sz = str(int(AUDIT_PT * 2))
            etree.SubElement(rPr, f"{{{W}}}sz").set(f"{{{W}}}val",   sz)
            etree.SubElement(rPr, f"{{{W}}}szCs").set(f"{{{W}}}val", sz)
            t = etree.SubElement(r, f"{{{W}}}t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = text
        # Label + TAB so the colon always starts at the tab stop
        _run(f"{label}\t", bold=True)
        _run(": ",          bold=True)
        _run(value or "-")

    # ── Indent Placed By ─────────────────────────────────────────────────────
    p_placed = _find_para(doc, "Indent Placed By")
    if p_placed:
        placed_by = getattr(indent, "indent_placed_by", None) or "-"
        _write_audit_para(p_placed, "Indent Placed By", placed_by)
        p_placed.paragraph_format.space_before = Pt(5)

        # ── Created By ───────────────────────────────────────────────────────────
    p_created = _find_para(doc, "Created")
    if p_created:
        val = f"{prepared_by['name']}  [{prepared_by['at']}]" if prepared_by else "-"
        _write_audit_para(p_created, "Created By", val)

    # ── Verified By (one row per checker; always at least one row) ───────────
    p_verified = _find_para(doc, "Verified By")
    if p_verified:
        if checked_by:
            _write_audit_para(p_verified, "Verified By",
                              f"{checked_by[0]['name']}  [{checked_by[0]['at']}]")
            prev_p = p_verified
            for extra in checked_by[1:]:
                new_p_el = _duplicate_para_after(prev_p)
                from docx.text.paragraph import Paragraph
                prev_p = Paragraph(new_p_el, doc)
                _write_audit_para(prev_p, "Verified By",
                                  f"{extra['name']}  [{extra['at']}]")
        else:
            _write_audit_para(p_verified, "Verified By", "-")

    # ── Approved By ──────────────────────────────────────────────────────────
    p_approved = _find_para(doc, "Approved By")
    if p_approved:
        val = (f"{authorised_by['name']}  [{authorised_by['at']}]"
               if authorised_by else "-")
        _write_audit_para(p_approved, "Approved By", val)


# ═════════════════════════════════════════════════════════════════════════════
# Items table helpers (low-level XML)
# ═════════════════════════════════════════════════════════════════════════════

def _merge_adjacent_tcs(tr_el, idx_a, idx_b):
    """Expand tc at idx_a by the gridSpan of tc at idx_b, then remove idx_b."""
    tcs = tr_el.findall(f"{{{W}}}tc")
    if len(tcs) <= max(idx_a, idx_b):
        return
    tc_a, tc_b = tcs[idx_a], tcs[idx_b]

    def _span(tc):
        tcPr = tc.find(f"{{{W}}}tcPr")
        if tcPr is not None:
            gs = tcPr.find(f"{{{W}}}gridSpan")
            if gs is not None:
                return int(gs.get(f"{{{W}}}val", "1"))
        return 1

    total = _span(tc_a) + _span(tc_b)
    tcPr_a = tc_a.find(f"{{{W}}}tcPr")
    if tcPr_a is None:
        tcPr_a = etree.Element(f"{{{W}}}tcPr")
        tc_a.insert(0, tcPr_a)
    for old in tcPr_a.findall(f"{{{W}}}gridSpan"):
        tcPr_a.remove(old)
    gs_el = etree.SubElement(tcPr_a, f"{{{W}}}gridSpan")
    gs_el.set(f"{{{W}}}val", str(total))
    tc_b.getparent().remove(tc_b)


def _apply_cant_split_tr(tr_el):
    """Add cantSplit to a <w:tr> element so the row never breaks across pages."""
    tr_pr = tr_el.find(f"{{{W}}}trPr")
    if tr_pr is None:
        tr_pr = etree.SubElement(tr_el, f"{{{W}}}trPr")
        tr_el.insert(0, tr_pr)
    for old in tr_pr.findall(f"{{{W}}}cantSplit"):
        tr_pr.remove(old)
    cs = etree.SubElement(tr_pr, f"{{{W}}}cantSplit")
    cs.set(f"{{{W}}}val", "1")


def _rpr_el(bold=False, size_pt=9, font="Bahnschrift Light"):
    """Return a fully-formed <w:rPr> lxml element."""
    rPr = etree.Element(f"{{{W}}}rPr")
    rFonts = etree.SubElement(rPr, f"{{{W}}}rFonts")
    rFonts.set(f"{{{W}}}ascii", font)
    rFonts.set(f"{{{W}}}hAnsi", font)
    if bold:
        etree.SubElement(rPr, f"{{{W}}}b")
    sz_val = str(int(size_pt * 2))
    etree.SubElement(rPr, f"{{{W}}}sz").set(f"{{{W}}}val",   sz_val)
    etree.SubElement(rPr, f"{{{W}}}szCs").set(f"{{{W}}}val", sz_val)
    return rPr


def _para_el(jc="left", sp_before=0, sp_after=0):
    """Return a <w:p> element with basic pPr (alignment + optional spacing)."""
    p = etree.Element(f"{{{W}}}p")
    pPr = etree.SubElement(p, f"{{{W}}}pPr")
    etree.SubElement(pPr, f"{{{W}}}jc").set(f"{{{W}}}val", jc)
    if sp_before or sp_after:
        sp = etree.SubElement(pPr, f"{{{W}}}spacing")
        if sp_before:
            sp.set(f"{{{W}}}before", str(sp_before))
        if sp_after:
            sp.set(f"{{{W}}}after",  str(sp_after))
    return p


def _tc_set_text(tc, text, jc="left", bold=False, size_pt=9,
                 font="Bahnschrift Light"):
    """Replace all paragraphs in tc with a single styled text paragraph."""
    for p in list(tc.findall(f"{{{W}}}p")):
        tc.remove(p)
    p = _para_el(jc=jc)
    r = etree.SubElement(p, f"{{{W}}}r")
    r.insert(0, _rpr_el(bold=bold, size_pt=size_pt, font=font))
    t = etree.SubElement(r, f"{{{W}}}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text or ""
    tc.append(p)


def _tc_set_item_desc(tc, item_name, note, name_pt=11, font="Bahnschrift Light"):
    """
    Fill a merged 'Item Name & Description' cell using a SINGLE paragraph:
      Run 1 — item_name  bold, name_pt
      Run 2 — <w:br/>    (soft line break, no extra paragraph spacing)
      Run 3 — note       non-bold, (name_pt − 2)pt
    sp_before=0, sp_after=0, single line-spacing on the paragraph.
    Cell vAlign (top) must be set by the caller.
    """
    for p in list(tc.findall(f"{{{W}}}p")):
        tc.remove(p)

    desc_pt = max(6, name_pt - 2)

    # One paragraph — left, 0/0 spacing, single line-height
    p = etree.Element(f"{{{W}}}p")
    pPr = etree.SubElement(p, f"{{{W}}}pPr")
    etree.SubElement(pPr, f"{{{W}}}jc").set(f"{{{W}}}val", "left")
    sp_el = etree.SubElement(pPr, f"{{{W}}}spacing")
    sp_el.set(f"{{{W}}}before",   "0")
    sp_el.set(f"{{{W}}}after",    "0")
    sp_el.set(f"{{{W}}}line",     "240")   # 240 twips = single line spacing
    sp_el.set(f"{{{W}}}lineRule", "auto")

    # Run 1: item name (bold)
    r1 = etree.SubElement(p, f"{{{W}}}r")
    r1.append(_rpr_el(bold=True, size_pt=name_pt, font=font))
    t1 = etree.SubElement(r1, f"{{{W}}}t")
    t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t1.text = item_name or ""

    if note:
        # Run 2: soft line break (carries desc font so line height is desc_pt)
        r_br = etree.SubElement(p, f"{{{W}}}r")
        r_br.append(_rpr_el(bold=False, size_pt=desc_pt, font=font))
        etree.SubElement(r_br, f"{{{W}}}br")

        # Run 3: description (smaller, normal weight)
        r2 = etree.SubElement(p, f"{{{W}}}r")
        r2.append(_rpr_el(bold=False, size_pt=desc_pt, font=font))
        t2 = etree.SubElement(r2, f"{{{W}}}t")
        t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t2.text = note

    tc.append(p)


# ═════════════════════════════════════════════════════════════════════════════
# Items table fill
# ═════════════════════════════════════════════════════════════════════════════

def _fill_table(doc, items):
    # Find by header text — index shifts after detail table insertion
    table = None
    for t in doc.tables:
        if t.rows and t.rows[0].cells[0].text.strip() == "SL No":
            table = t
            break
    if table is None:
        return

    FONT = "Bahnschrift Light"
    # After merging cols 2+3: SL No=0, Item Code=1, Item Name & Desc=2, Unit=3, Qty=4, Location=5
    CENTER_COLS = {0, 1, 3, 4}
    HDR_TEXTS   = ["SL No", "Item Code", "Item Name & Description",
                   "Unit", "Qty", "Location"]

    # ── Header row: merge cols 2+3, restyle all cells ────────────────────────
    hdr_tr = table.rows[0]._tr
    _merge_adjacent_tcs(hdr_tr, 2, 3)
    _apply_cant_split_tr(hdr_tr)

    # Repeat header row on every page when table spans multiple pages
    hdr_trPr = hdr_tr.find(f"{{{W}}}trPr")
    if hdr_trPr is None:
        hdr_trPr = etree.SubElement(hdr_tr, f"{{{W}}}trPr")
        hdr_tr.insert(0, hdr_trPr)
    etree.SubElement(hdr_trPr, f"{{{W}}}tblHeader")

    for ci, tc in enumerate(hdr_tr.findall(f"{{{W}}}tc")):
        tcPr = tc.find(f"{{{W}}}tcPr")
        if tcPr is None:
            tcPr = etree.Element(f"{{{W}}}tcPr")
            tc.insert(0, tcPr)
        for old in tcPr.findall(f"{{{W}}}vAlign"):
            tcPr.remove(old)
        etree.SubElement(tcPr, f"{{{W}}}vAlign").set(f"{{{W}}}val", "center")

        jc_val = "center" if ci in CENTER_COLS else "left"
        _tc_set_text(tc, HDR_TEXTS[ci] if ci < len(HDR_TEXTS) else "",
                     jc=jc_val, bold=True, size_pt=12, font=FONT)

    # ── Template row: merge cols 2+3 ─────────────────────────────────────────
    template_row = copy.deepcopy(table.rows[1]._tr)
    _merge_adjacent_tcs(template_row, 2, 3)

    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    min_rows = max(6, len(items))

    # ── Data rows ─────────────────────────────────────────────────────────────
    def _build_data_row(item):
        new_tr = copy.deepcopy(template_row)
        _apply_cant_split_tr(new_tr)
        qty_str = f"{item['qty']:.2f}" if item["qty"] else ""
        simple = {
            0: str(item["sl"]) if item["sl"] else "",
            1: item["item_code"],
            3: item["unit"],
            4: qty_str,
            5: item["location"],
        }
        for ci, tc in enumerate(new_tr.findall(f"{{{W}}}tc")):
            tcPr = tc.find(f"{{{W}}}tcPr")
            if tcPr is None:
                tcPr = etree.Element(f"{{{W}}}tcPr")
                tc.insert(0, tcPr)
            for old in tcPr.findall(f"{{{W}}}vAlign"):
                tcPr.remove(old)
            va_val = "top" if ci == 2 else "center"
            etree.SubElement(tcPr, f"{{{W}}}vAlign").set(f"{{{W}}}val", va_val)
            for old in tcPr.findall(f"{{{W}}}noWrap"):
                tcPr.remove(old)

            if ci == 2:
                _tc_set_item_desc(tc, item["item_name"], item["note"],
                                  name_pt=12, font=FONT)
            elif ci == 5:                          # Location — smaller font
                _tc_set_text(tc, simple.get(5, ""),
                             jc="left", bold=False, size_pt=9, font=FONT)
            else:
                val    = simple.get(ci, "")
                jc_val = "center" if ci in CENTER_COLS else "left"
                _tc_set_text(tc, val, jc=jc_val, bold=False, size_pt=12, font=FONT)
        return new_tr

    for item in items:
        table._tbl.append(_build_data_row(item))

    # ── Empty filler rows ─────────────────────────────────────────────────────
    empty_item = {"sl": "", "item_code": "", "item_name": "", "note": "",
                  "unit": "", "qty": 0, "location": ""}
    for _ in range(min_rows - len(items)):
        table._tbl.append(_build_data_row({**empty_item, "sl": ""}))

    # ── Blue footer row at end of table ──────────────────────────────────────
    blue_tr = copy.deepcopy(template_row)
    for tc in blue_tr.findall(f"{{{W}}}tc"):
        tcPr = tc.find(f"{{{W}}}tcPr")
        if tcPr is None:
            tcPr = etree.Element(f"{{{W}}}tcPr")
            tc.insert(0, tcPr)
        # Remove existing shd, add blue fill
        for old in tcPr.findall(f"{{{W}}}shd"):
            tcPr.remove(old)
        shd = etree.SubElement(tcPr, f"{{{W}}}shd")
        shd.set(f"{{{W}}}val", "clear")
        shd.set(f"{{{W}}}color", "auto")
        shd.set(f"{{{W}}}fill", "B6DDE8")  # blue hex — change this to adjust shade
        # Clear any text in this row
        for p in tc.findall(f"{{{W}}}p"):
            for r in p.findall(f"{{{W}}}r"):
                p.remove(r)
    table._tbl.append(blue_tr)


# ═════════════════════════════════════════════════════════════════════════════
# Public API
# ═════════════════════════════════════════════════════════════════════════════

def generate_indent_pdf(indent_id, base_url, force=False):
    try:
        indent = IndentMaster.query.get(indent_id)
        if not indent:
            return res("Indent not found", [], 404)

        # Serve cached PDF when within expiry (bypass with ?force=1)
        if (not force
                and getattr(indent, "pdf_url",          None)
                and getattr(indent, "pdf_generated_at", None)):
            age = (datetime.utcnow() - indent.pdf_generated_at).days
            if age < PDF_EXPIRY_DAYS:
                base_prefix = current_app.config.get(
                    "INDENT_PDF_BASE_URL", "/resource/indent/pdf-file")
                relative = indent.pdf_url.replace(base_prefix + "/", "", 1)
                return serve_indent_pdf(relative)

        # ── build core data ───────────────────────────────────────────────────
        fp           = _fingerprint(indent)
        generated_at = datetime.utcnow().strftime("%d-%b-%Y %H:%M UTC")
        token        = _serializer().dumps(
            {"iid": indent.id, "ino": indent.indent_no, "fp": fp},
            salt=PDF_TOKEN_SALT,
        )
        verify_url = f"{base_url.rstrip('/')}/resource/indent/verify/{token}"

        items = [
            {
                "sl":        idx,
                "item_code": ii.item_code or "",
                "item_name": (ii.item.item_name if ii.item else "") or "",
                "note":      ii.note or "",
                "unit":      (ii.item.unit.unit_name
                              if ii.item and ii.item.unit else "") or "",
                "qty":       float(ii.qty or 0),
                "location":  getattr(ii, "location", "") or "",
            }
            for idx, ii in enumerate(indent.indent_items, 1)
        ]

        prepared_by, checked_by, authorised_by = _build_signatures(indent.id)
        is_approved = authorised_by is not None

        # ── build document ────────────────────────────────────────────────────
        doc = Document(TEMPLATE_PATH)

        # Fill text fields first
        _fill_doc(doc, indent, prepared_by, checked_by, authorised_by)
        _fill_table(doc, items)

        # Repeating header: logo + website + QR (also clears body P0 floating images)
        qr_bytes = _qr_with_brackets(verify_url)
        _build_header(doc, qr_bytes, generated_at)

        # Hash line in body (tamper evidence on the paper itself)
        _add_hash_line(doc, fp, generated_at)

        # Micro-print footer (every page) + generated timestamp
        _add_security_footer(doc, indent.indent_no, fp, generated_at)

        if is_approved:
            # APPROVED diagonal watermark in header → repeats on every page
            wm_el = _watermark_para_el("APPROVED", color_hex="D8D8D8")
            doc.sections[0].header._element.append(wm_el)

            # Green tick + VERIFIED text + VERIFIED watermark on last page
            # _add_verified_section(doc, fp)

        # ── docx → PDF via MS Word COM ────────────────────────────────────────
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "indent.docx")
            pdf_path  = os.path.join(tmpdir, "indent.pdf")
            doc.save(docx_path)
            pythoncom.CoInitialize()
            try:
                convert(docx_path, pdf_path)
            finally:
                pythoncom.CoUninitialize()
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

        pdf_url = _save_pdf(pdf_bytes, indent.indent_no)

        if hasattr(indent, "pdf_url"):
            indent.pdf_url          = pdf_url
            indent.pdf_token        = token
            indent.pdf_generated_at = datetime.utcnow()
            db.session.commit()

        response = make_response(pdf_bytes)
        response.headers["Content-Type"]        = "application/pdf"
        response.headers["Content-Disposition"] = (
            f'inline; filename="{indent.indent_no}.pdf"'
        )
        return response

    except Exception as e:
        db.session.rollback()
        return res(str(e), [], 500)


def verify_indent_pdf(token):
    try:
        data   = _serializer().loads(token, salt=PDF_TOKEN_SALT,
                                     max_age=PDF_EXPIRY_DAYS * 86400)
        indent = IndentMaster.query.get(data["iid"])

        if not indent or getattr(indent, "pdf_token", None) != token:
            return _verify_page("invalid")

        live_fp    = _fingerprint(indent)
        stored_fp  = data.get("fp", "")

        if stored_fp != live_fp:
            return _verify_page(
                "tampered",
                indent_no  = data.get("ino", ""),
                stored_fp  = stored_fp,
                live_fp    = live_fp,
            )

        return _verify_page(
            "valid",
            indent_no = indent.indent_no,
            fp        = live_fp,
            project   = indent.project_code or "",
        )

    except SignatureExpired as e:
        try:
            payload = e.payload
            if payload:
                indent = IndentMaster.query.get(payload.get("iid"))
                if indent and getattr(indent, "pdf_token", None) == token:
                    indent.pdf_url = indent.pdf_token = indent.pdf_generated_at = None
                    db.session.commit()
        except Exception:
            pass
        return _verify_page("expired")

    except BadSignature:
        return _verify_page("invalid")

    except Exception as e:
        return _verify_page("error", msg=str(e))


def _verify_page(status, **kw):
    color_map = {
        "valid":    "#1b5e20",
        "tampered": "#b71c1c",
        "expired":  "#e65100",
        "invalid":  "#37474f",
        "error":    "#37474f",
    }
    title_map = {
        "valid":    "✓ INDENT VERIFIED",
        "tampered": "⚠ TAMPERED DOCUMENT",
        "expired":  "⏱ QR CODE EXPIRED",
        "invalid":  "✗ INVALID QR",
        "error":    "Error",
    }
    color = color_map.get(status, "#37474f")
    title = title_map.get(status, "Unknown")

    if status == "valid":
        body = f"""
          <p style="font-size:48px;color:#2e7d32;margin:0;">✔</p>
          <p>This indent is <strong>authentic and unmodified</strong>.</p>
          <table style="margin:14px auto;font-size:13px;border-collapse:collapse;
                        text-align:left;">
            <tr><td style="padding:3px 14px;font-weight:bold;">Indent No</td>
                <td>: {kw.get('indent_no','')}</td></tr>
            <tr><td style="padding:3px 14px;font-weight:bold;">Site Code</td>
                <td>: {kw.get('project','')}</td></tr>
            <tr><td style="padding:3px 14px;font-weight:bold;">Document Hash</td>
                <td style="font-family:monospace;font-size:10px;">
                  : {kw.get('fp','')}</td></tr>
          </table>
          <p style="font-size:11px;color:#888;">
            Hash matches server record — document has not been altered.</p>"""
        code = 200

    elif status == "tampered":
        body = f"""
          <p>Data mismatch — <strong>do not accept this document.</strong></p>
          <table style="margin:14px auto;font-size:12px;border-collapse:collapse;
                        text-align:left;">
            <tr><td style="padding:3px 14px;font-weight:bold;">Indent</td>
                <td>: {kw.get('indent_no','')}</td></tr>
            <tr><td style="padding:3px 14px;font-weight:bold;">PDF Hash</td>
                <td style="font-family:monospace;font-size:10px;">
                  : {kw.get('stored_fp','')}</td></tr>
            <tr><td style="padding:3px 14px;font-weight:bold;">Live Hash</td>
                <td style="font-family:monospace;font-size:10px;">
                  : {kw.get('live_fp','')}</td></tr>
          </table>"""
        code = 400

    elif status == "expired":
        body = ("<p>This QR has expired (valid 2 days only).</p>"
                "<p>Request a freshly generated PDF.</p>")
        code = 410

    elif status == "invalid":
        body = "<p>This QR code is not recognised by the system.</p>"
        code = 400

    else:
        body = f"<p>{kw.get('msg', 'Unknown error')}</p>"
        code = 500

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>Indent Verification — Dishaan Hi-Tech</title>
<style>
  body{{font-family:Arial,sans-serif;display:flex;justify-content:center;
       align-items:center;min-height:100vh;margin:0;background:#f5f5f5;}}
  .card{{background:#fff;border-radius:12px;padding:40px 48px;max-width:560px;
         width:90%;text-align:center;box-shadow:0 4px 20px rgba(0,0,0,.12);}}
  .title{{font-size:20px;font-weight:bold;color:{color};margin:8px 0 16px;}}
  .body{{font-size:14px;color:#444;line-height:1.8;}}
  .footer{{font-size:10px;color:#bbb;margin-top:24px;}}
</style></head><body>
<div class="card">
  <div class="title">{title}</div>
  <div class="body">{body}</div>
  <div class="footer">Dishaan Hi-Tech (India) Private Limited</div>
</div></body></html>"""
    return make_response(html, code)
