
import os
import io
import tempfile
import subprocess

from datetime import datetime,date
from sqlalchemy.exc import SQLAlchemyError

from app.extensions import db
from app.response import res
from app.models.project import Project
from app.models.category_group import CategoryMaster
from app.models.item import Item
from app.modules.work_flow import *

from app.models.indent_master import IndentMaster
from app.models.indent_item import IndentItem
from app.models.approval_path import *
from app.cloudinary_uploader import *



# GENERATE INDENT NUMBER


def generate_indent_no(project_code):

    last_indent = (
        db.session.query(
            IndentMaster.indent_no
        )
        .filter(
            IndentMaster.project_code == project_code
        )
        .order_by(
            IndentMaster.id.desc()
        )
        .first()
    )

    if last_indent:
        try:
            last_serial = int(
                last_indent[0]
            )
        except:
            last_serial = 350000
    else:
        last_serial = 350000

    return str(last_serial + 1)


# CREATE INDENT


def create_indent(data,files=None, created_by=None):

    try:

        project = Project.query.filter_by(
            project_code=data.get("projectCode")
        ).first()

        if not project:
            return res("Invalid project code", [], 400)

        # ==========================================
        # VALIDATE CATEGORY
        # ==========================================

        category = CategoryMaster.query.filter_by(
            fixed_code=data.get("categoryCode")
        ).first()

        if not category:
            return res("Invalid category code", [], 400)


        # VALIDATE ITEMS


        items = data.get("items", [])

        if not items:
            return res("Indent items required", [], 400)


        xtemp=generate_indent_no(
                data.get("projectCode")
            )
        supporting_file = None

        indent_file = files.get(
            "indentFile"
        )

        if not indent_file:
            return res(
                "Indent file is required",
                [],
                400
            )
        supporting_file = (

                upload_file_to_bunny(

                    file=
                    indent_file,

                    mainFolder=
                    "indent",

                    subFolder=
                   xtemp,

                    fileName=
                    "support"

                )

            )
        # CREATE MASTER

        indent = IndentMaster(

            indent_no=xtemp,
            supporting_file= supporting_file,
            project_code=data.get(
                "projectCode"
            ),

            category_code=data.get(
                "categoryCode"
            ),

            indent_date=data.get("indentDate"),

            priority=data.get(
                "priority"
            ),

            required_within=data.get(
                "requiredWithin"
            ),

            indent_placed_by=data.get(
                "indentPlacedBy"
            ),

            site_reg_serial_no=data.get(
                "siteRegSerialNo"
            ),

            sale_order_no=data.get(
                "saleOrderNo"
            ),

            remarks=data.get(
                "remarks"
            ),

            order_status="Pending",

            # workflow

            workflow_status="Draft",

            current_level=0,

            locked=False,

            created_by=created_by
        )

        db.session.add(indent)
        db.session.flush()

        # ==========================================
        # CREATE ITEMS
        # ==========================================

        for row in items:

            item = Item.query.filter_by(
                item_code=row.get("itemCode")
            ).first()

            if not item:
                db.session.rollback()
                return res(
                    f"Invalid item code : {row.get('itemCode')}",
                    [],
                    400
                )

            # ======================================
            # CATEGORY MATCH VALIDATION
            # ======================================

            if item.category_code != data.get("categoryCode"):
                db.session.rollback()

                return res(
                    f"Item {item.item_code} "
                    f"does not belong to category "
                    f"{data.get('categoryCode')}",
                    [],
                    400
                )

            # ======================================
            # QTY VALIDATION
            # ======================================

            qty = row.get("qty")

            if not qty or float(qty) <= 0:
                db.session.rollback()

                return res(
                    f"Invalid qty for item "
                    f"{item.item_code}",
                    [],
                    400
                )

            indent_item = IndentItem(

                indent_id=indent.id,

                item_code=item.item_code,

                qty=qty,

                location=row.get("location"),

                note=row.get("note"),

                created_by=created_by
            )

            db.session.add(indent_item)

        # COMMIT


        db.session.commit()

        return res(
            "Indent created successfully",
            {
                "indentId": indent.id,
                "indentNo": indent.indent_no
            },
            200
        )

    except SQLAlchemyError as e:

        db.session.rollback()

        return res(str(e), [], 500)

    except Exception as e:

        db.session.rollback()

        return res(str(e), [], 500)


# INDENT LIST


def get_indent_list(filters=None):

    try:

        query = IndentMaster.query

        if filters:

            if filters.get("projectCode"):
                query = query.filter(
                    IndentMaster.project_code ==
                    filters.get("projectCode")
                )

            if filters.get("status"):
                query = query.filter(
                    IndentMaster.workflow_status ==
                    filters.get("status")
                )

            if filters.get("categoryCode"):
                query = query.filter(
                    IndentMaster.category_code ==
                    filters.get("categoryCode")
                )

        indents = (
            query
            .order_by(IndentMaster.id.desc())
            .all()
        )

        data = []

        for row in indents:

            data.append({

                "id": row.id,

                "indentNo": row.indent_no,

                "projectCode": row.project_code,

                "categoryCode": row.category_code,

                "indentCategoryName": row.category.category_name,

                "priority": row.priority,

                "indentStatus": row.workflow_status,

                "orderStatus": row.order_status,
                "indentDate" : row.indent_date,
                "placedBy" : row.creator.username if row.creator else None,
                "createdAt": row.created_at.strftime(
                    "%Y-%m-%d %H:%M:%S"
                ) if row.created_at else None
            })

        return res(
            "Indent list fetched successfully",
            data,
            200
        )

    except Exception as e:

        return res(str(e), [], 500)


# =========================================================
# GET SINGLE INDENT
# =========================================================

def get_indent_details(indent_id):

    try:

        indent = IndentMaster.query.get(indent_id)

        if not indent:
            return res("Indent not found", [], 404)

        item_rows = []

        for row in indent.indent_items:

            item_rows.append({

                "id": row.id,

                "itemCode": row.item_code,

                "itemName": row.item.item_name
                if row.item else None,

                "qty": float(row.qty),

                "location": row.location,

                "note": row.note,

                "unit": (
                    row.item.unit.short_name
                    if row.item and row.item.unit
                    else None
                )
            })

        response = {

            "id": indent.id,

            "indentNo": indent.indent_no,

            "projectCode": indent.project_code,

            "categoryCode": indent.category_code,
            "indentDate": str(indent.indent_date) if indent.indent_date else None,
            "priority": indent.priority,
            "siteRegSerialNo":indent.site_reg_serial_no,
            "saleOrderNo":indent.sale_order_no,
            "requiredWithin": str(indent.required_within)
            if indent.required_within else None,

            "indentPlacedBy": indent.indent_placed_by,

            "remarks": indent.remarks,

            "indentStatus": indent.workflow_status,

            "items": item_rows,
            "currentLevel": indent.current_level,
            "indentFile":indent.supporting_file,
            "locked": indent.locked,
            "submittedAt":
                str(indent.submitted_at)
                if indent.submitted_at
                else None,
            "finalApprovedAt":
                str(indent.final_approved_at)
                if indent.final_approved_at
                else None,
            "rejectedAt":
                str(indent.rejected_at)
                if indent.rejected_at
                else None,
        }

        return res(
            "Indent details fetched successfully",
            response,
            200
        )

    except Exception as e:

        return res(str(e), [], 500)

def get_items_by_category(category_code):

    try:

        items = Item.query.filter_by(
            category_code=category_code,
            status="Active"
        ).all()

        result = []

        for row in items:

            result.append({

                "itemCode": row.item_code,

                "itemName": row.item_name,

                "description": row.item_description,

                "unit": (
                    row.unit.short_name
                    if row.unit else None
                ),

                "gst": (
                    float(row.gst_percentage)
                    if row.gst_percentage else 0
                )
            })

        return res(
            "Items fetched successfully",
            result,
            200
        )

    except Exception as e:

        return res(str(e), [], 500)

def update_indent(indent_id, data, files=None, updated_by=None):

    try:

        indent = IndentMaster.query.get(indent_id)

        if not indent:
            return res("Indent not found", [], 404)

        # ==========================================
        # ALLOW ONLY DRAFT EDIT
        # ==========================================

        if indent.locked:
            return res("Indent cannot be edited", [], 400)

        # ==========================================
        # REMARKS REQUIRED
        # ==========================================

        remarks = data.get("remarks")

        if not remarks:
            return res("Remarks required", [], 400)

        indent.remarks = remarks

        # ==========================================
        # CATEGORY CHANGE OPTIONAL
        # ==========================================

        if "categoryCode" in data:

            category = CategoryMaster.query.filter_by(fixed_code=data.get("categoryCode")).first()

            if not category:
                return res("Invalid category code", [], 400)

            indent.category_code = data.get("categoryCode")

        category_code = indent.category_code

        # ==========================================
        # UPDATE OPTIONAL FIELDS
        # ==========================================

        if "priority" in data:
            indent.priority = data["priority"]

        if "requiredWithin" in data:
            indent.required_within = data["requiredWithin"]

        if "indentPlacedBy" in data:
            indent.indent_placed_by = data["indentPlacedBy"]

        if "siteRegSerialNo" in data:
            indent.site_reg_serial_no = data["siteRegSerialNo"]

        if "saleOrderNo" in data:
            indent.sale_order_no = data["saleOrderNo"]

        # ==========================================
        # UPDATE FILE OPTIONAL
        # ==========================================

        if files:

            indent_file = files.get(
                "indentFile"
            )

            if indent_file:
                indent.supporting_file = (

                    upload_file_to_bunny(

                        file=
                        indent_file,

                        mainFolder=
                        "indent",

                        subFolder=
                        indent.indent_no ,

                        fileName=
                        "support"

                    )

                )
        # ==========================================
        # UPDATE ITEMS ONLY IF SENT
        # ==========================================

        items = data.get("items")

        if items is not None:

            IndentItem.query.filter_by(indent_id=indent.id).delete()

            for row in items:

                item = Item.query.filter_by(item_code=row.get("itemCode")).first()

                if not item:
                    db.session.rollback()
                    return res(f"Invalid item code: {row.get('itemCode')}", [], 400)

                if item.category_code != category_code:
                    db.session.rollback()
                    return res(f"Item {item.item_code} does not belong to category {category_code}", [], 400)

                qty = row.get("qty")

                if not qty or float(qty) <= 0:
                    return res(f"Invalid qty for item {item.item_code}", [], 400)

                indent_item = IndentItem(
                    indent_id=indent.id,
                    item_code=item.item_code,
                    qty=qty,
                    location=row.get("location"),
                    note=row.get("note"),
                    created_by=updated_by
                )

                db.session.add(indent_item)

        # ==========================================

        if indent.workflow_status == "Reback":
            indent.correction_sent_at = None

        indent.updated_by = updated_by
        indent.updated_at = datetime.utcnow()

        db.session.commit()

        return res("Indent updated successfully", {
            "indentId": indent.id,
            "indentNo": indent.indent_no
        }, 200)

    except SQLAlchemyError as e:

        db.session.rollback()
        return res(str(e), [], 500)

    except Exception as e:

        db.session.rollback()
        return res(str(e), [], 500)

# SUBMIT INDENT


def submit_indent(
        indent_id,
        submitted_by=None
):

    try:

        indent = IndentMaster.query.get(
            indent_id
        )

        if not indent:
            return res(
                "Indent not found",
                [],
                404
            )

        # ==========================================
        # ALREADY SUBMITTED
        # ==========================================

        if indent.workflow_status not in [
            "Draft",
            "Reback"
        ]:

            return res(
                "Indent already submitted",
                [],
                400
            )

        # ==========================================
        # VALIDATE ITEMS EXIST
        # ==========================================

        if not indent.indent_items:

            return res(
                "Indent has no items",
                [],
                400
            )

        # ==========================================
        # RESTART AFTER REBACK
        # ==========================================

        if indent.workflow_status == "Reback":

            indent.current_level = 0

        # ==========================================
        # FIND FIRST APPROVER
        # ==========================================

        first_level = get_first_approver(

            indent.project_code,

            "indent"
        )

        # ==========================================
        # NO APPROVER
        # AUTO APPROVE
        # ==========================================

        if not first_level:

            indent.workflow_status = (
                "Approved"
            )

            indent.locked = True

            indent.approved_by = (
                submitted_by
            )

            indent.submitted_at = (
                datetime.utcnow() )

            indent.final_approved_at = (
                datetime.utcnow()
            )

        else:

            # ======================================
            # START APPROVAL FLOW
            # ======================================

            indent.workflow_status = (
                f"Pending_L"
                f"{first_level.level_no}"
            )

            indent.current_level = (
                first_level.level_no
            )

            indent.locked = True

            indent.submitted_at = (
                datetime.utcnow()
            )


        # HISTORY
        #

        create_history(

            project_code=
            indent.project_code,

            module_code=
            "indent",

            record_id=
            indent.id,

            level_no=
            indent.current_level,

            action=
            "SUBMIT",

            action_by=
            submitted_by
        )

        indent.updated_by = (
            submitted_by
        )
        indent.submitted_by = (submitted_by)

        indent.updated_at = (
            datetime.utcnow()
        )

        db.session.commit()

        return res(
            "Indent submitted successfully",
            {
                "indentId":
                    indent.id,

                "indentNo":
                    indent.indent_no,

                "workflowStatus":
                    indent.workflow_status
            },
            200
        )

    except SQLAlchemyError as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )

    except Exception as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )

# =========================================================
# DELETE DRAFT INDENT
# =========================================================

def delete_indent(indent_id):

    try:

        indent = IndentMaster.query.get(indent_id)

        if not indent:
            return res("Indent not found", [], 404)

        # ==========================================
        # ALLOW ONLY DRAFT DELETE
        # ==========================================

        if indent.locked:
            return res(
                "Only editable indent can be deleted",
                [],
                400
            )

        # ==========================================
        # DELETE ITEMS
        # ==========================================

        IndentItem.query.filter_by(
            indent_id=indent.id
        ).delete()

        # ==========================================
        # DELETE MASTER
        # ==========================================

        db.session.delete(indent)

        db.session.commit()

        return res(
            "Indent deleted successfully",
            [],
            200
        )

    except SQLAlchemyError as e:

        db.session.rollback()

        return res(str(e), [], 500)

    except Exception as e:

        db.session.rollback()

        return res(str(e), [], 500)



def approve_indent(
        indent_id,
        approved_by=None,
        comments=None
):

    try:

        indent = IndentMaster.query.get(
            indent_id
        )

        if not indent:

            return res(
                "Indent not found",
                [],
                404
            )

        # ==========================================
        # ONLY PENDING CAN APPROVE
        # ==========================================

        if not indent.workflow_status.startswith(
            "Pending"
        ):

            return res(
                "Indent not pending",
                [],
                400
            )

        # ==========================================
        # CHECK CURRENT APPROVER
        # ==========================================

        allowed = is_current_approver(

            indent.project_code,

            "indent",

            indent.current_level,

            approved_by
        )
        print(
            "project:",
            indent.project_code
        )

        print(
            "module:",
            "indent"
        )

        print(
            "level:",
            indent.current_level,
            type(
                indent.current_level
            )
        )

        print(
            "approved_by:",
            approved_by,
            type(
                approved_by
            )
        )
        if not allowed:

            return res(
                "You are not current approver",
                [],
                403
            )

        # ==========================================
        # FIND NEXT LEVEL
        # ==========================================

        next_level = get_next_approver(

            indent.project_code,

            "indent",

            indent.current_level
        )

        # ==========================================
        # NEXT LEVEL EXISTS
        # ==========================================

        if next_level:

            create_history(

                project_code=
                indent.project_code,

                module_code=
                "indent",

                record_id=
                indent.id,

                level_no=
                indent.current_level,

                action=
                "APPROVE",

                action_by=
                approved_by,

                comments=
                comments
            )

            indent.current_level = (
                next_level.level_no
            )

            indent.workflow_status = (

                f"Pending_L"
                f"{next_level.level_no}"

            )

        else:

            # ======================================
            # FINAL APPROVE
            # ======================================

            create_history(

                project_code=
                indent.project_code,

                module_code=
                "indent",

                record_id=
                indent.id,

                level_no=
                indent.current_level,

                action=
                "FINAL_APPROVE",

                action_by=
                approved_by,

                comments=
                comments
            )

            indent.workflow_status = (
                "Approved"
            )

            indent.locked = True

            indent.approved_by = (
                approved_by
            )

            indent.final_approved_at = (
                datetime.utcnow()
            )

        indent.updated_by = (
            approved_by
        )

        indent.updated_at = (
            datetime.utcnow()
        )

        db.session.commit()

        return res(

            "Indent approved successfully",

            {

                "indentId":
                indent.id,

                "workflowStatus":
                indent.workflow_status,

                "currentLevel":
                indent.current_level

            },

            200
        )

    except SQLAlchemyError as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )

    except Exception as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )


def reback_indent(
        indent_id,
        reback_by=None,
        comments=None
):

    try:

        indent = IndentMaster.query.get(
            indent_id
        )

        if not indent:

            return res(
                "Indent not found",
                [],
                404
            )

        # ==========================================
        # ONLY PENDING
        # ==========================================

        if not indent.workflow_status.startswith(
            "Pending"
        ):

            return res(
                "Indent not pending",
                [],
                400
            )

        # ==========================================
        # COMMENT REQUIRED
        # ==========================================

        if not comments:

            return res(
                "Comments required",
                [],
                400
            )

        # ==========================================
        # CURRENT APPROVER CHECK
        # ==========================================

        allowed = is_current_approver(

            indent.project_code,

            "indent",

            indent.current_level,

            reback_by
        )

        if not allowed:

            return res(
                "You are not current approver",
                [],
                403
            )

        # ==========================================
        # UPDATE STATUS
        # ==========================================

        indent.workflow_status = (
            "Reback"
        )

        indent.locked = False

        indent.correction_sent_at = (
            datetime.utcnow()
        )

        indent.updated_by = (
            reback_by
        )

        indent.updated_at = (
            datetime.utcnow()
        )

        # ==========================================
        # HISTORY
        # ==========================================

        create_history(

            project_code=
            indent.project_code,

            module_code=
            "indent",

            record_id=
            indent.id,

            level_no=
            indent.current_level,

            action=
            "REBACK",

            action_by=
            reback_by,

            comments=
            comments
        )

        db.session.commit()

        return res(

            "Indent sent for correction",

            {

                "indentId":
                indent.id,

                "workflowStatus":
                indent.workflow_status

            },

            200
        )

    except SQLAlchemyError as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )

    except Exception as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )


def reject_indent(
        indent_id,
        rejected_by=None,
        comments=None
):

    try:

        indent = IndentMaster.query.get(
            indent_id
        )

        if not indent:

            return res(
                "Indent not found",
                [],
                404
            )

        # ==========================================
        # ONLY PENDING
        # ==========================================

        if not indent.workflow_status.startswith(
            "Pending"
        ):

            return res(
                "Indent not pending",
                [],
                400
            )

        # ==========================================
        # COMMENT REQUIRED
        # ==========================================

        if not comments:

            return res(
                "Comments required",
                [],
                400
            )

        # ==========================================
        # CURRENT APPROVER CHECK
        # ==========================================

        allowed = is_current_approver(

            indent.project_code,

            "indent",

            indent.current_level,

            rejected_by
        )

        if not allowed:

            return res(
                "You are not current approver",
                [],
                403
            )

        # ==========================================
        # REJECT
        # ==========================================

        indent.workflow_status = (
            "Rejected"
        )

        indent.locked = True

        indent.rejected_at = (
            datetime.utcnow()
        )

        indent.rejected_by = (
            rejected_by
        )

        indent.status = (
            "Inactive"
        )

        indent.updated_by = (
            rejected_by
        )

        indent.updated_at = (
            datetime.utcnow()
        )

        # ==========================================
        # HISTORY
        # ==========================================

        create_history(

            project_code=
            indent.project_code,

            module_code=
            "indent",

            record_id=
            indent.id,

            level_no=
            indent.current_level,

            action=
            "REJECT",

            action_by=
            rejected_by,

            comments=
            comments
        )

        db.session.add(
            indent
        )

        db.session.commit()

        return res(

            "Indent rejected successfully",

            {

                "indentId":
                indent.id,

                "workflowStatus":
                indent.workflow_status

            },

            200
        )

    except SQLAlchemyError as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )

    except Exception as e:

        db.session.rollback()

        return res(
            str(e),
            [],
            500
        )


def get_indent_history(
        indent_id
):

    try:

        indent = IndentMaster.query.get(
            indent_id
        )

        if not indent:

            return res(
                "Indent not found",
                [],
                404
            )

        rows = get_history(

            "indent",

            indent.id
        )

        data=[]

        for row in rows:

            data.append({

                "id":
                row.id,

                "action":
                row.action,

                "level":
                row.level_no,

                "comments":
                row.comments,

                "actionBy":
                (
                    row.user.username
                    if row.user
                    else None
                ),

                "createdAt":
                row.created_at.strftime(
                    "%Y-%m-%d %H:%M:%S"
                )
                if row.created_at
                else None
            })

        return res(

            "Indent history fetched successfully",

            data,

            200
        )

    except Exception as e:

        return res(
            str(e),
            [],
            500
        )


# =========================================================
# GENERATE INDENT DOCUMENT  (fills Indent.docx template)
# =========================================================
#
# Requirements:
#   pip install docxtpl
#
# Template setup (one-time):
#   Open  asset/Indent.docx  in Word and replace every
#   hardcoded placeholder value with the Jinja2 tags shown
#   in the context dict below.  For the items table add a
#   {% for row in items %} / {% endfor %} loop row.
#
#   Key tags to place inside the .docx:
#
#   Header block
#   ─────────────────────────────────────────────────────
#   {{site_code}}           ← Site Code cell
#   {{indent_no}}           ← Indent No cell
#   {{project_name}}        ← Project Name cell
#   {{indent_date}}         ← Indent Date cell
#   {{customer_name}}       ← Customer Name cell
#   {{required_date}}       ← Required Date cell
#   {{sale_order_no}}       ← Sale Order No cell
#   {{indent_category}}     ← Indent Category cell
#
#   Items table (one template row, wrapped in for-loop)
#   ─────────────────────────────────────────────────────
#   {% for row in items %}
#   {{row.sl_no}}  {{row.item_code}}  {{row.item_name}}
#   {{row.specification}}  {{row.unit}}  {{row.qty}}
#   {{row.location}}
#   {% endfor %}
#
#   Signature block
#   ─────────────────────────────────────────────────────
#   {{indent_placed_by}}
#   {{created_by}}   {{created_at}}
#   {{verified_by}}  {{verified_at}}
#   {{approved_by}}  {{approved_at}}
# =========================================================

# Path to the Word template (relative to project root)
_INDENT_TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "../../../../asset/indent.docx"
)


class _FileWrapper:
    """
    Wraps a bytes payload so upload_file_to_bunny()
    can call  .filename  and  .read()  on it, exactly
    as it would on a Werkzeug FileStorage object.
    """

    def __init__(self, data: bytes, filename: str):
        self.filename = filename
        self._data = data

    def read(self) -> bytes:
        return self._data


def _add_cell_borders(cell) -> None:
    """Apply thin black borders to all 4 sides of a python-docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), "000000")
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _fmt_dt(dt_str: str | None, fmt: str = "%d-%b-%Y") -> str:
    """Format a datetime/date string for display, e.g. '07-May-2026'."""
    if not dt_str:
        return "—"
    for pattern in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(dt_str, pattern).strftime(fmt)
        except ValueError:
            pass
    return dt_str  # fallback: return as-is


def generate_indent_document(indent_id) -> tuple:
    """
    Generates a filled Indent DOCX from  asset/Indent.docx,
    uploads it to Bunny CDN, and returns the public URL.

    Usage (in a route / controller):
        response, status = generate_indent_document(indent_id)

    Returns:
        res("...", {"documentUrl": "<cdn_url>"}, 200)  on success
        res("...", [], 4xx/5xx)                        on failure
    """
    try:
        from docxtpl import DocxTemplate

        # ── 1. Indent details ────────────────────────────────────
        detail_payload, detail_code = get_indent_details(indent_id)

        if detail_code != 200:
            return res(
                detail_payload.get("message", "Indent not found"),
                [],
                detail_code
            )

        detail = detail_payload["data"]

        # ── 2. Indent workflow history ───────────────────────────
        history_payload, history_code = get_indent_history(indent_id)

        history = (
            history_payload["data"]
            if history_code == 200
            else []
        )

        # ── 3. Project info (name + customer) ───────────────────
        project = Project.query.filter_by(
            project_code=detail.get("projectCode")
        ).first()

        project_name  = project.project_name  if project else "—"
        customer_name = project.client_name   if project else "—"

        # ── 4. Signature data from workflow history ──────────────
        #
        # History actions produced by the workflow:
        #   SUBMIT        → who submitted (created_by in the doc)
        #   APPROVE       → intermediate approver  (Verified By)
        #   FINAL_APPROVE → last approver           (Approved By)
        #   REBACK        → correction              (skip)
        #   REJECT        → rejection               (skip)

        created_by  = "—";  created_at  = "—"
        verified_by = "—";  verified_at = "—"
        approved_by = "—";  approved_at = "—"

        for h in history:

            action = (h.get("action") or "").upper()

            actor = h.get("actionBy") or "—"
            ts    = _fmt_dt(
                h.get("createdAt"),
                fmt="%d-%b-%Y; %H:%M"
            )

            if action == "SUBMIT":
                created_by = actor
                created_at = ts

            elif action == "APPROVE":
                # First (intermediate) approval = Verified By
                if verified_by == "—":
                    verified_by = actor
                    verified_at = ts

            elif action == "FINAL_APPROVE":
                approved_by = actor
                approved_at = ts

        # ── 5. Items rows ────────────────────────────────────────
        items = []

        for idx, row in enumerate(detail.get("items", []), start=1):
            items.append({
                "sl_no":         idx,
                "item_code":     row.get("itemCode")  or "—",
                "item_name":     row.get("itemName")  or "—",
                "specification": row.get("note")      or "—",
                "unit":          row.get("unit")      or "—",
                "qty":           row.get("qty")       or 0,
                "location":      row.get("location")  or "—",
            })

        # ── 6. Build Jinja2 context (no items — added via python-docx below) ──
        context = {

            # Header
            "site_code":       detail.get("siteRegSerialNo") or "—",
            "indent_no":       detail.get("indentNo")        or "—",
            "project_name":    project_name,
            "indent_date":     _fmt_dt(detail.get("indentDate")),
            "customer_name":   customer_name,
            "required_date":   _fmt_dt(detail.get("requiredWithin")),
            "sale_order_no":   detail.get("saleOrderNo")     or "—",
            "indent_category": detail.get("categoryCode")    or "—",

            # Signature block
            "indent_placed_by": detail.get("indentPlacedBy") or "—",
            "created_by":       created_by,
            "created_at":       created_at,
            "verified_by":      verified_by,
            "verified_at":      verified_at,
            "approved_by":      approved_by,
            "approved_at":      approved_at,
        }

        # ── 7. Render template with docxtpl ─────────────────────
        tpl = DocxTemplate(_INDENT_TEMPLATE_PATH)
        tpl.render(context)

        indent_no_safe = (
            str(detail.get("indentNo", indent_id))
            .replace("/", "_")
            .replace(" ", "_")
        )

        # Save rendered DOCX to a temp file
        tmp_dir   = tempfile.mkdtemp()
        docx_path = os.path.join(tmp_dir, f"indent_{indent_no_safe}.docx")
        pdf_path  = os.path.join(tmp_dir, f"indent_{indent_no_safe}.pdf")

        tpl.save(docx_path)

        # ── 7b. Add items rows via python-docx ──────────────────
        #
        #   The items table is the 3rd table in the document (index 2):
        #   [0] banner  [1] info-header  [2] items  [3] signature
        #
        from docx import Document as _DocxDoc
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        _doc = _DocxDoc(docx_path)
        itbl = _doc.tables[2]   # items table

        _COL_ALIGN = [
            WD_ALIGN_PARAGRAPH.CENTER,  # Sl No
            WD_ALIGN_PARAGRAPH.CENTER,  # Item Code
            WD_ALIGN_PARAGRAPH.LEFT,    # Item Name
            WD_ALIGN_PARAGRAPH.LEFT,    # Specification (note)
            WD_ALIGN_PARAGRAPH.CENTER,  # Unit
            WD_ALIGN_PARAGRAPH.CENTER,  # Qty
            WD_ALIGN_PARAGRAPH.LEFT,    # Location
        ]

        for item in items:
            row_data = [
                str(item["sl_no"]),
                str(item["item_code"]),
                str(item["item_name"]),
                str(item["specification"]),   # ← note field
                str(item["unit"]),
                str(item["qty"]),
                str(item["location"]),
            ]
            new_row = itbl.add_row()
            for cell, text, align in zip(new_row.cells, row_data, _COL_ALIGN):
                para = cell.paragraphs[0]
                para.clear()
                run  = para.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(9)
                para.alignment = align
                _add_cell_borders(cell)

        _doc.save(docx_path)

        # ── 8. Convert DOCX → PDF ────────────────────────────────
        #
        #   Windows (dev)  : uses docx2pdf  →  pip install docx2pdf
        #                    (requires Microsoft Word to be installed)
        #
        #   Linux (Railway): uses LibreOffice headless
        #                    Add to Railway nixpacks / Dockerfile:
        #                    apt-get install -y libreoffice
        #
        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(docx_path, pdf_path)

        except ImportError:
            # Fallback: LibreOffice headless (Linux / Railway)
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmp_dir,
                    docx_path,
                ],
                check=True,
                timeout=60
            )

        if not os.path.exists(pdf_path):
            return res(
                "PDF conversion failed — "
                "install docx2pdf (Windows) or libreoffice (Linux/Railway)",
                [],
                500
            )

        # ── 9. Upload PDF to Bunny CDN ───────────────────────────
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        # Clean up temp files
        try:
            os.unlink(docx_path)
            os.unlink(pdf_path)
            os.rmdir(tmp_dir)
        except Exception:
            pass

        file_wrapper = _FileWrapper(
            data=pdf_bytes,
            filename=f"indent_{indent_no_safe}.pdf"
        )

        print (file_wrapper)
        cdn_url = upload_file_to_bunny(
            file=file_wrapper,
            mainFolder="indent_docs",
            subFolder=detail.get("projectCode", "unknown"),
            fileName=f"indent_{indent_no_safe}"
        )

        if not cdn_url:
            return res(
                "PDF generated but upload to CDN failed",
                [],
                500
            )

        return res(
            "Indent document generated successfully",
            {"documentUrl": cdn_url},
            200
        )

    except ImportError:
        return res(
            "docxtpl is not installed. "
            "Run: pip install docxtpl",
            [],
            500
        )

    except Exception as e:
        return res(str(e), [], 500)