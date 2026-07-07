"""
Order PDF service — exact replication of Order.docx template.

Template specs (extracted from document.xml):
  Page    : US Letter 8.5"×11", portrait
  Margins : top=1843, bottom=709, left=567, right=758, header=708 (all twips)
  Font    : Aptos Display 9pt (body), Segoe UI 9pt (header company info),
            Aptos Display 16pt bold (title), Aptos 10pt (GSTIN/summaries)
  Colors  : D9D9D9 (gray headers), BFBFBF (final total), 7F7F7F (sub-desc text)
  Borders : single sz=4 (0.5pt) black — all sides (TableGrid style)
  Grid    : 10-column table, col widths (twips):
            562, 2935, 972, 715, 913, 964, 1135, 583, 991, 1135 = 10905 total
"""
import io, os, copy, hashlib, tempfile, subprocess
from datetime import datetime
from lxml import etree
from PIL import Image, ImageDraw
import qrcode
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature
from docx import Document
from docx.shared import Inches, Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from flask import current_app, make_response, send_from_directory

from app.extensions import db
from app.models.orderMaster import OrderMaster
from app.models.companies import Company
from app.models.approval_path import ApprovalHistory
from app.models.user import User
from app.response import res

# ── XML namespace ─────────────────────────────────────────────────────────────
W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
W10  = "urn:schemas-microsoft-com:office:word"

# ── Page geometry (twips) ─────────────────────────────────────────────────────
_PG_W  = 12240   # 8.5"
_PG_H  = 15840   # 11"
_MG_T  = 1843    # top
_MG_B  = 709     # bottom
_MG_L  = 567     # left
_MG_R  = 758     # right
_MG_HD = 708     # header distance

# ── 10-column grid (twips) ────────────────────────────────────────────────────
# Cols: SL | ItemDetails | HSN | Unit | Qty | Rate | BasicAmt | GST% | GSTAmt | TotalAmt
COL_W = [562, 2935, 972, 715, 913, 964, 1135, 583, 991, 1135]
# Merged column spans for parties block:
#   Vendor  = cols 0+1   = 562+2935        = 3497
#   Company = cols 2-5   = 972+715+913+964 = 3564
#   Ref     = cols 6-9   = 1135+583+991+1135 = 3844

# ── Typography ────────────────────────────────────────────────────────────────
F_BODY  = "Aptos Display"   # 9pt — all table cells, body text
F_HDR   = "Segoe UI"        # 9pt — company info in letterhead
F_TITLE = "Aptos Display"   # 16pt bold — "PURCHASES ORDER"
F_GSTIN = "Aptos"           # 10pt — GSTIN line

# ── Colors ────────────────────────────────────────────────────────────────────
C_GRAY_HD  = "D9D9D9"   # header rows (light gray)
C_GRAY_TOT = "BFBFBF"   # final total row (darker gray)
C_GRAY_TXT = "7F7F7F"   # sub-description text

# ── Token / storage ───────────────────────────────────────────────────────────
PDF_TOKEN_SALT  = "order-pdf-v2"
PDF_EXPIRY_DAYS = 2
LOGO_PATH = os.getenv("COMPANY_LOGO_PATH",
                       os.path.join(os.getcwd(), "asset", "order", "1000018063.jpeg"))


# ═════════════════════════════════════════════════════════════════════════════
# Storage helpers
# ═════════════════════════════════════════════════════════════════════════════

def _storage_root():
    return current_app.config.get(
        "PDF_STORAGE_PATH",
        os.path.join(os.getcwd(), "storage", "pdf"),
    )


def _save_pdf(pdf_bytes, order_no):
    root   = _storage_root()
    folder = os.path.join(root, "order_pdf", order_no)
    os.makedirs(folder, exist_ok=True)
    with open(os.path.join(folder, "order.pdf"), "wb") as f:
        f.write(pdf_bytes)
    base = current_app.config.get("PDF_BASE_URL", "/resource/order/pdf-file")
    return f"{base}/order_pdf/{order_no}/order.pdf"


def serve_pdf_file(relative_path):
    root     = _storage_root()
    full_dir = os.path.join(root, os.path.dirname(relative_path))
    return send_from_directory(full_dir, os.path.basename(relative_path),
                               mimetype="application/pdf")


# ═════════════════════════════════════════════════════════════════════════════
# Crypto / fingerprint
# ═════════════════════════════════════════════════════════════════════════════

def _serializer():
    secret = current_app.config.get("JWT_SECRET_KEY", "erp-secret")
    return URLSafeTimedSerializer(secret)


def _fingerprint(order):
    raw = f"{order.order_no}|{float(order.basic_amount or 0):.2f}|{float(order.total_amount or 0):.2f}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16].upper()


# ═════════════════════════════════════════════════════════════════════════════
# QR with L-bracket corners
# ═════════════════════════════════════════════════════════════════════════════

def _qr_with_brackets(url):
    from qrcode.image.pil import PilImage
    qr = qrcode.QRCode(box_size=10, border=4,
                       error_correction=qrcode.constants.ERROR_CORRECT_M)
    qr.add_data(url)
    qr.make(fit=True)
    buf = io.BytesIO()
    qr.make_image(image_factory=PilImage,
                  fill_color="black", back_color="white").save(buf, "PNG")
    buf.seek(0)
    img  = Image.open(buf).convert("RGB")
    draw = ImageDraw.Draw(img)
    iw, ih = img.size
    t = max(5, iw // 45)
    L = iw // 5
    for (x0, y0, x1, y1) in [
        (0,    0,    L,    t), (0,    0,    t,    L),
        (iw-L, 0,    iw,   t), (iw-t, 0,    iw,   L),
        (0,    ih-t, L,    ih), (0,    ih-L, t,    ih),
        (iw-L, ih-t, iw,   ih), (iw-t, ih-L, iw,   ih),
    ]:
        draw.rectangle([x0, y0, x1-1, y1-1], fill="black")
    out = io.BytesIO()
    img.save(out, "PNG")
    return out.getvalue()


# ═════════════════════════════════════════════════════════════════════════════
# Low-level XML helpers
# ═════════════════════════════════════════════════════════════════════════════

def _w(tag):
    return f"{{{W}}}{tag}"


def _set_table_grid(table, col_widths_twips):
    """Set exact column widths on the tblGrid element."""
    tbl = table._tbl
    for old in tbl.findall(_w("tblGrid")):
        tbl.remove(old)
    # Build element first, then insert at correct position (after tblPr, before rows)
    grid = etree.Element(_w("tblGrid"))
    for w in col_widths_twips:
        col = etree.SubElement(grid, _w("gridCol"))
        col.set(_w("w"), str(w))
    tblPr = tbl.find(_w("tblPr"))
    if tblPr is not None:
        tbl.insert(list(tbl).index(tblPr) + 1, grid)
    else:
        tbl.insert(0, grid)


def _set_table_width(table, width_twips):
    tbl   = table._tbl
    tblPr = tbl.find(_w("tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, _w("tblPr"))
    for old in tblPr.findall(_w("tblW")):
        tblPr.remove(old)
    tblW = etree.SubElement(tblPr, _w("tblW"))
    tblW.set(_w("w"), str(width_twips))
    tblW.set(_w("type"), "dxa")


def _set_table_borders(table, sz=4, color="auto"):
    """Apply TableGrid-style borders: all sides single."""
    tbl   = table._tbl
    tblPr = tbl.find(_w("tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, _w("tblPr"))
    for old in tblPr.findall(_w("tblBorders")):
        tblPr.remove(old)
    borders = etree.SubElement(tblPr, _w("tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = etree.SubElement(borders, _w(side))
        el.set(_w("val"),   "single")
        el.set(_w("sz"),    str(sz))
        el.set(_w("space"), "0")
        el.set(_w("color"), color)


def _set_no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(_w("tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, _w("tblPr"))
    for old in tblPr.findall(_w("tblBorders")):
        tblPr.remove(old)
    borders = etree.SubElement(tblPr, _w("tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = etree.SubElement(borders, _w(side))
        el.set(_w("val"),   "none")
        el.set(_w("sz"),    "0")
        el.set(_w("space"), "0")
        el.set(_w("color"), "auto")


def _cell_shading(cell, fill):
    tcPr = cell._tc.find(_w("tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(cell._tc, _w("tcPr"))
        cell._tc.insert(0, tcPr)
    for old in tcPr.findall(_w("shd")):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, _w("shd"))
    shd.set(_w("val"),   "clear")
    shd.set(_w("color"), "auto")
    shd.set(_w("fill"),  fill)


def _cell_width(cell, width_twips):
    tcPr = cell._tc.find(_w("tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(cell._tc, _w("tcPr"))
        cell._tc.insert(0, tcPr)
    for old in tcPr.findall(_w("tcW")):
        tcPr.remove(old)
    tcW = etree.SubElement(tcPr, _w("tcW"))
    tcW.set(_w("w"),    str(width_twips))
    tcW.set(_w("type"), "dxa")


def _cell_valign(cell, val="center"):
    tcPr = cell._tc.find(_w("tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(cell._tc, _w("tcPr"))
        cell._tc.insert(0, tcPr)
    for old in tcPr.findall(_w("vAlign")):
        tcPr.remove(old)
    etree.SubElement(tcPr, _w("vAlign")).set(_w("val"), val)


def _cell_margins(cell, top=0, left=108, bottom=0, right=108):
    """Set cell internal margins (twips). Template default: 0/108/0/108."""
    tcPr = cell._tc.find(_w("tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(cell._tc, _w("tcPr"))
        cell._tc.insert(0, tcPr)
    for old in tcPr.findall(_w("tcMar")):
        tcPr.remove(old)
    mar = etree.SubElement(tcPr, _w("tcMar"))
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = etree.SubElement(mar, _w(side))
        el.set(_w("w"),    str(val))
        el.set(_w("type"), "dxa")


def _row_height(row, height_twips, exact=False):
    trPr = row._tr.find(_w("trPr"))
    if trPr is None:
        trPr = etree.SubElement(row._tr, _w("trPr"))
        row._tr.insert(0, trPr)
    for old in trPr.findall(_w("trHeight")):
        trPr.remove(old)
    h = etree.SubElement(trPr, _w("trHeight"))
    h.set(_w("val"),   str(height_twips))
    if exact:
        h.set(_w("hRule"), "exact")


def _cant_split(row):
    trPr = row._tr.find(_w("trPr"))
    if trPr is None:
        trPr = etree.SubElement(row._tr, _w("trPr"))
        row._tr.insert(0, trPr)
    for old in trPr.findall(_w("cantSplit")):
        trPr.remove(old)
    etree.SubElement(trPr, _w("cantSplit")).set(_w("val"), "1")


def _repeat_header(row):
    """Mark a row as a repeating table header."""
    trPr = row._tr.find(_w("trPr"))
    if trPr is None:
        trPr = etree.SubElement(row._tr, _w("trPr"))
        row._tr.insert(0, trPr)
    if not trPr.findall(_w("tblHeader")):
        etree.SubElement(trPr, _w("tblHeader"))


def _para(cell, text, font=F_BODY, size_pt=9, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
          sp_before=0, sp_after=0, line=240):
    """Clear the cell and write one paragraph with one run."""
    tc = cell._tc
    for p_el in tc.findall(_w("p")):
        tc.remove(p_el)
    p_el = etree.SubElement(tc, _w("p"))

    # pPr
    pPr = etree.SubElement(p_el, _w("pPr"))
    jc_map = {
        WD_ALIGN_PARAGRAPH.LEFT:    "left",
        WD_ALIGN_PARAGRAPH.CENTER:  "center",
        WD_ALIGN_PARAGRAPH.RIGHT:   "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
    }
    etree.SubElement(pPr, _w("jc")).set(_w("val"), jc_map.get(align, "left"))
    sp = etree.SubElement(pPr, _w("spacing"))
    sp.set(_w("before"), str(int(sp_before * 20)))
    sp.set(_w("after"),  str(int(sp_after  * 20)))
    sp.set(_w("line"),   str(line))
    sp.set(_w("lineRule"), "auto")

    # run
    r_el = etree.SubElement(p_el, _w("r"))
    rPr  = etree.SubElement(r_el, _w("rPr"))
    rf   = etree.SubElement(rPr, _w("rFonts"))
    rf.set(_w("ascii"), font); rf.set(_w("hAnsi"), font); rf.set(_w("cs"), font)
    if bold:    etree.SubElement(rPr, _w("b"))
    if italic:  etree.SubElement(rPr, _w("i"))
    sz_val = str(int(size_pt * 2))
    etree.SubElement(rPr, _w("sz")).set(_w("val"),   sz_val)
    etree.SubElement(rPr, _w("szCs")).set(_w("val"), sz_val)
    if color:
        etree.SubElement(rPr, _w("color")).set(_w("val"), color)
    t = etree.SubElement(r_el, _w("t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = str(text) if text is not None else ""
    return p_el


def _para_multiline(cell, lines, font=F_BODY, size_pt=9, bold=False,
                    align=WD_ALIGN_PARAGRAPH.LEFT, sp_before=0, sp_after=0):
    """Write multiple paragraphs into a cell (one per line)."""
    tc = cell._tc
    for p_el in tc.findall(_w("p")):
        tc.remove(p_el)
    for i, (text, b, sz, col) in enumerate(lines):
        p_el = etree.SubElement(tc, _w("p"))
        pPr  = etree.SubElement(p_el, _w("pPr"))
        etree.SubElement(pPr, _w("jc")).set(_w("val"),
            "center" if align == WD_ALIGN_PARAGRAPH.CENTER else
            "right"  if align == WD_ALIGN_PARAGRAPH.RIGHT  else "left")
        sp = etree.SubElement(pPr, _w("spacing"))
        sp.set(_w("before"), str(int((sp_before if i == 0 else 0) * 20)))
        sp.set(_w("after"),  str(int((sp_after  if i == len(lines)-1 else 0) * 20)))
        sp.set(_w("line"),   "240"); sp.set(_w("lineRule"), "auto")

        if text is not None and str(text).strip() != "":
            r_el = etree.SubElement(p_el, _w("r"))
            rPr  = etree.SubElement(r_el, _w("rPr"))
            rf   = etree.SubElement(rPr, _w("rFonts"))
            rf.set(_w("ascii"), font); rf.set(_w("hAnsi"), font); rf.set(_w("cs"), font)
            if b: etree.SubElement(rPr, _w("b"))
            sz_val = str(int(sz * 2))
            etree.SubElement(rPr, _w("sz")).set(_w("val"),   sz_val)
            etree.SubElement(rPr, _w("szCs")).set(_w("val"), sz_val)
            if col:
                etree.SubElement(rPr, _w("color")).set(_w("val"), col)
            t = etree.SubElement(r_el, _w("t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = str(text)


def _horiz_merge(row, start_col, end_col):
    """Merge cells [start_col..end_col] in a row (0-indexed). Returns merged cell."""
    cells = row.cells
    c = cells[start_col].merge(cells[end_col])
    return c


# ═════════════════════════════════════════════════════════════════════════════
# Word section header: Logo | PURCHASE ORDER | QR
# ═════════════════════════════════════════════════════════════════════════════

def _build_word_header(doc, qr_bytes, company):
    sec    = doc.sections[0]
    sec.header_distance = Twips(_MG_HD)
    header = sec.header
    header.is_linked_to_previous = False

    for p in list(header.paragraphs):
        p._p.getparent().remove(p._p)

    # 3-column borderless header table
    ht = header.add_table(rows=1, cols=3, width=Twips(sum(COL_W)))
    _set_no_borders(ht)
    ht.autofit = False; ht.allow_autofit = False

    # col widths: Logo=2.43" | Title=2.71" | QR=2.43"
    # roughly matching parties col widths: 3497 | 3411 | 3997
    hdr_col_w = [3497, 3411, 3997]
    cells = ht.rows[0].cells
    for ci, w in enumerate(hdr_col_w):
        _cell_width(cells[ci], w)

    # ── Left: Logo ────────────────────────────────────────────────────────────
    lp = cells[0].paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(0)
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(1.3))

    # Company name + address in header left (Segoe UI 9pt)
    def _hdr_para(cell, text, bold=False):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        r.font.name = F_HDR
        r.font.size = Pt(9)
        r.font.bold = bold

    _hdr_para(cells[0], company.company_name or "", bold=True)
    _hdr_para(cells[0], company.registered_address or "")
    _hdr_para(cells[0],
              f"Ph: {company.contact_number or ''}  |  Web: www.dishaanhitech.com")

    # ── Center: Title + website ───────────────────────────────────────────────
    tp = cells[1].paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(8)
    tp.paragraph_format.space_after  = Pt(2)
    tr = tp.add_run("PURCHASES ORDER")
    tr.font.name = F_TITLE; tr.font.size = Pt(16); tr.font.bold = True

    wp = cells[1].add_paragraph()
    wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wp.paragraph_format.space_before = Pt(0)
    wp.paragraph_format.space_after  = Pt(0)
    wr = wp.add_run("www.dishaanhitech.com")
    wr.font.name = F_HDR; wr.font.size = Pt(9); wr.font.bold = True

    # ── Right: QR ────────────────────────────────────────────────────────────
    qp = cells[2].paragraphs[0]
    qp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    qp.paragraph_format.space_before = Pt(0)
    qp.paragraph_format.space_after  = Pt(0)
    qp.add_run().add_picture(io.BytesIO(qr_bytes), width=Inches(1.2))


# ═════════════════════════════════════════════════════════════════════════════
# Word footer: micro-print generated line
# ═════════════════════════════════════════════════════════════════════════════

def _build_word_footer(doc, generated_at, fp):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._p.getparent().remove(p._p)
    fp_p = footer.add_paragraph()
    fp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp_p.paragraph_format.space_before = Pt(0)
    fp_p.paragraph_format.space_after  = Pt(0)
    r = fp_p.add_run(
        f"Generated: {generated_at}  |  Fingerprint: {fp}  |  "
        f"Dishaan Hi-Tech (India) Pvt. Ltd.  |  www.dishaanhitech.com"
    )
    r.font.name = F_HDR; r.font.size = Pt(4)


# ═════════════════════════════════════════════════════════════════════════════
# Main data table (single 10-column table, all merges applied)
# ═════════════════════════════════════════════════════════════════════════════

def _build_main_table(doc, order, vendor, company, items):
    n_items = max(len(items), 4)   # at least 4 item rows like template

    # Total rows:
    #   0 = parties header (D9D9D9)
    #   1 = parties content
    #   2 = preamble (merged)
    #   3 = items header (D9D9D9)
    #   4..4+n_items-1 = item rows
    #   4+n_items = TOTAL row (D9D9D9)
    #   4+n_items+1 = blank / basic amount label row
    #   4+n_items+2..5 = CGST, SGST, IGST (3 rows)
    #   4+n_items+6 = Total Order Amount (BFBFBF)
    #   4+n_items+7 = Amount in words (full merge)

    R_PARTY_HDR  = 0
    R_PARTY_BODY = 1
    R_PREAMBLE   = 2
    R_ITEMS_HDR  = 3
    R_ITEMS_START = 4
    R_ITEMS_END   = 3 + n_items
    R_TOTAL       = 4 + n_items
    R_BASIC       = 5 + n_items
    R_CGST        = 6 + n_items
    R_SGST        = 7 + n_items
    R_IGST        = 8 + n_items
    R_GRAND       = 9 + n_items
    R_WORDS       = 10 + n_items
    n_rows        = 11 + n_items

    tbl = doc.add_table(rows=n_rows, cols=10)
    tbl.autofit = False; tbl.allow_autofit = False
    _set_table_borders(tbl)
    _set_table_width(tbl, sum(COL_W))
    _set_table_grid(tbl, COL_W)

    # Set default cell widths for all rows
    for row in tbl.rows:
        for ci, cw in enumerate(COL_W):
            _cell_width(row.cells[ci], cw)
            _cell_margins(row.cells[ci], top=0, left=108, bottom=0, right=108)

    # ── Row 0: Parties header ─────────────────────────────────────────────────
    r0 = tbl.rows[R_PARTY_HDR]
    _row_height(r0, 283)
    _cant_split(r0)
    _r0 = list(r0.cells)   # pre-store all 10 refs before any merge
    c_vendor = _r0[0].merge(_r0[1])   # cols 0-1 = 3497
    c_comp   = _r0[2].merge(_r0[5])   # cols 2-5 = 3564
    c_ref    = _r0[6].merge(_r0[9])   # cols 6-9 = 3844
    for c in (c_vendor, c_comp, c_ref):
        _cell_shading(c, C_GRAY_HD)
        _cell_valign(c, "center")
    _para(c_vendor, "Service From (Saler):", font=F_BODY, size_pt=9, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(c_comp,   "Service To / Buyer:",   font=F_BODY, size_pt=9, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(c_ref,    "Reference:",             font=F_BODY, size_pt=9, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_width(c_vendor, 3497)
    _cell_width(c_comp,   3564)
    _cell_width(c_ref,    3844)

    # ── Row 1: Parties content ────────────────────────────────────────────────
    r1 = tbl.rows[R_PARTY_BODY]
    _cant_split(r1)
    _r1 = list(r1.cells)   # pre-store all 10 refs before any merge
    cv = _r1[0].merge(_r1[1])
    cc = _r1[2].merge(_r1[5])
    cr = _r1[6].merge(_r1[9])
    _cell_valign(cv, "top"); _cell_valign(cc, "top"); _cell_valign(cr, "top")
    _cell_width(cv, 3497); _cell_width(cc, 3564); _cell_width(cr, 3844)

    v = vendor
    _para_multiline(cv, [
        (v.ledger_name or "—",            True,  9, None),
        (v.registered_address or "",      False, 9, None),
        (f"Email: {getattr(v,'email','') or '—'}",    False, 9, None),
        (f"Contact: {v.primary_contact_person or '—'}", False, 9, None),
        (f"Ph: {v.primary_contact_number or '—'}",      False, 9, None),
        (f"PAN: {v.pan or '—'}",          False, 9, None),
        (f"GSTIN: {v.gstin or '—'}",      False, 9, None),
        (f"State: {v.state_name or '—'}   Code: {v.state_code or '—'}", False, 9, None),
    ])

    c = company
    _para_multiline(cc, [
        (c.company_name or "—",           True,  9, None),
        (c.registered_address or "",      False, 9, None),
        (f"Email: {c.email or '—'}",      False, 9, None),
        (f"Contact: {c.contact_person or '—'}", False, 9, None),
        (f"Ph: {c.contact_number or '—'}",     False, 9, None),
        (f"PAN: {c.pan or '—'}",          False, 9, None),
        (f"GSTIN: {c.gstn or '—'}",       False, 9, None),
        (f"State: {c.state or '—'}   Code: {c.state_code or '—'}", False, 9, None),
    ])

    _para_multiline(cr, [
        (f"Order No    :  {order.order_no or '—'}",         True,  9, None),
        (f"Order Date  :  {order.order_date or '—'}",       False, 9, None),
        (f"Quotation No:  {order.quotation_no or '—'}",     False, 9, None),
        (f"Quot. Date  :  {order.quotation_date or '—'}",   False, 9, None),
        (f"Project     :  {order.project_code or '—'}",     False, 9, None),
        (f"Status      :  {order.workflow_status or '—'}",  False, 9, None),
    ])

    # ── Row 2: Preamble ───────────────────────────────────────────────────────
    r2 = tbl.rows[R_PREAMBLE]
    _r2 = list(r2.cells)
    cp = _r2[0].merge(_r2[9])
    _cell_width(cp, sum(COL_W))
    _cell_valign(cp, "top")
    _para_multiline(cp, [
        ("Dear Sir / Madam,", False, 9, None),
        ("With reference to your quotation and our discussions, we are pleased to place "
         "this Purchase Order for the supply of materials / services as detailed below. "
         "Please acknowledge receipt and confirm acceptance of this order.", False, 9, None),
    ], sp_before=3, sp_after=3)

    # ── Row 3: Items header ───────────────────────────────────────────────────
    r3 = tbl.rows[R_ITEMS_HDR]
    _row_height(r3, 340)
    _repeat_header(r3)
    _cant_split(r3)
    HDR_LABELS = ["SL.\nNo", "Item Details", "HSN/\nSAC", "Unit",
                  "Order\nQty", "Unit Rate\n(INR)", "Basic\nAmount",
                  "GST\n%", "GST\nAmount", "Total\nAmount"]
    for ci, cell in enumerate(r3.cells):
        _cell_shading(cell, C_GRAY_HD)
        _cell_valign(cell, "center")
        _para(cell, HDR_LABELS[ci], font=F_BODY, size_pt=9, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Rows 4..R_ITEMS_END: Item data ───────────────────────────────────────
    CENTER_COLS = {0, 2, 3, 4}
    RIGHT_COLS  = {5, 6, 8, 9}

    for ri in range(n_items):
        row = tbl.rows[R_ITEMS_START + ri]
        _cant_split(row)
        if ri < len(items):
            item = items[ri]
            vals = [
                str(item.get("sl", "")),
                "",                                         # Item Details handled below
                item.get("hsn_sac", "") or "",
                item.get("unit", "") or "",
                f"{item.get('qty', 0):.2f}",
                f"{item.get('rate', 0):,.2f}",
                f"{item.get('amount', 0):,.2f}",
                f"{item.get('gst_percent', 0):.1f}",
                f"{item.get('gst_amount', 0):,.2f}",
                f"{item.get('total', 0):,.2f}",
            ]
            for ci, cell in enumerate(row.cells):
                _cell_valign(cell, "top" if ci == 1 else "center")
                if ci == 1:
                    # Item name + sub-description (gray smaller text)
                    name = item.get("item_name", "") or ""
                    note = item.get("note", "") or ""
                    lines = [(name, True, 9, None)]
                    if note:
                        lines.append((note, False, 7, C_GRAY_TXT))
                    _para_multiline(cell, lines)
                else:
                    align = (WD_ALIGN_PARAGRAPH.CENTER if ci in CENTER_COLS else
                             WD_ALIGN_PARAGRAPH.RIGHT  if ci in RIGHT_COLS  else
                             WD_ALIGN_PARAGRAPH.LEFT)
                    _para(cell, vals[ci], font=F_BODY, size_pt=9, align=align)
        else:
            # Empty filler row
            for ci, cell in enumerate(row.cells):
                _cell_valign(cell, "center")
                _para(cell, "", font=F_BODY, size_pt=9)

    # ── TOTAL row ─────────────────────────────────────────────────────────────
    r_tot = tbl.rows[R_TOTAL]
    _row_height(r_tot, 283)
    _cant_split(r_tot)
    # Pre-store refs BEFORE merge so indices don't shift
    _tot_cells = list(r_tot.cells)
    ct_sl    = _tot_cells[0]
    ct_basic = _tot_cells[6]
    ct_gstpc = _tot_cells[7]
    ct_gstam = _tot_cells[8]
    ct_total = _tot_cells[9]
    # Merge cols 1-5 for "TOTAL" label: 2935+972+715+913+964 = 6499
    ct_label = _tot_cells[1].merge(_tot_cells[5])

    basic_total = sum(i.get("amount", 0) for i in items)
    gst_total   = sum(i.get("gst_amount", 0) for i in items)
    grand_total  = float(order.total_amount or 0)

    for c in (ct_sl, ct_label, ct_basic, ct_gstpc, ct_gstam, ct_total):
        _cell_shading(c, C_GRAY_HD)
        _cell_valign(c, "center")
    _para(ct_sl,    "",              font=F_BODY, size_pt=9)
    _para(ct_label, "TOTAL",         font=F_BODY, size_pt=9, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(ct_basic, f"{basic_total:,.2f}", font=F_BODY, size_pt=9, bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(ct_gstpc, "",              font=F_BODY, size_pt=9)
    _para(ct_gstam, f"{gst_total:,.2f}",   font=F_BODY, size_pt=9, bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(ct_total, f"{grand_total:,.2f}",  font=F_BODY, size_pt=9, bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_width(ct_sl,    COL_W[0])
    _cell_width(ct_label, sum(COL_W[1:6]))
    _cell_width(ct_basic, COL_W[6])
    _cell_width(ct_gstpc, COL_W[7])
    _cell_width(ct_gstam, COL_W[8])
    _cell_width(ct_total, COL_W[9])

    # ── Row Basic total (2-col: label left, amount right) ────────────────────
    r_basic = tbl.rows[R_BASIC]
    _cant_split(r_basic)
    _bc = list(r_basic.cells)   # pre-store before merges
    cb_label = _bc[0].merge(_bc[4])   # cols 0-4 = 6097
    cb_val   = _bc[5].merge(_bc[9])   # cols 5-9 = 4808
    _cell_valign(cb_label, "center"); _cell_valign(cb_val, "center")
    _cell_width(cb_label, sum(COL_W[0:5])); _cell_width(cb_val, sum(COL_W[5:10]))
    _para(cb_label, "Total Basic Amount",
          font=F_BODY, size_pt=10, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(cb_val, f"INR  {float(order.basic_amount or 0):,.2f}",
          font=F_BODY, size_pt=10, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Compute CGST/SGST/IGST (simple 50/50 split of gst_amount per item)
    cgst = sgst = igst = 0.0
    for oi in order.items:
        g = float(oi.gst_amount or 0)
        gp = float(oi.gst_percent or 0)
        cgst += g / 2
        sgst += g / 2

    # ── CGST / SGST / IGST rows (3-col each) ─────────────────────────────────
    tax_rows = [
        (R_CGST, "CGST @", cgst),
        (R_SGST, "SGST @", sgst),
        (R_IGST, "IGST @", igst),
    ]
    for r_idx, label, amt in tax_rows:
        row = tbl.rows[r_idx]
        _cant_split(row)
        _rc = list(row.cells)   # pre-store before any merges
        c0 = _rc[0].merge(_rc[4])    # cols 0-4 = 6097
        c1 = _rc[5].merge(_rc[7])    # cols 5-7 = 2682
        c2 = _rc[8].merge(_rc[9])    # cols 8-9 = 2126
        _cell_valign(c0, "center"); _cell_valign(c1, "center"); _cell_valign(c2, "center")
        _cell_width(c0, sum(COL_W[0:5]))
        _cell_width(c1, sum(COL_W[5:8]))
        _cell_width(c2, sum(COL_W[8:10]))
        _para(c0, label,              font=F_BODY, size_pt=10, bold=False,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
        _para(c1, "",                 font=F_BODY, size_pt=10)
        _para(c2, f"INR  {amt:,.2f}", font=F_BODY, size_pt=10, bold=False,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Grand total row (BFBFBF) ──────────────────────────────────────────────
    r_grand = tbl.rows[R_GRAND]
    _cant_split(r_grand)
    _gc = list(r_grand.cells)   # pre-store before merges
    cg0 = _gc[0].merge(_gc[4])
    cg1 = _gc[5].merge(_gc[7])
    cg2 = _gc[8].merge(_gc[9])
    for c in (cg0, cg1, cg2):
        _cell_shading(c, C_GRAY_TOT)
        _cell_valign(c, "center")
    _cell_width(cg0, sum(COL_W[0:5]))
    _cell_width(cg1, sum(COL_W[5:8]))
    _cell_width(cg2, sum(COL_W[8:10]))
    _para(cg0, "Total Order Amount",
          font=F_BODY, size_pt=10, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(cg1, "",               font=F_BODY, size_pt=10)
    _para(cg2, f"INR  {grand_total:,.2f}",
          font=F_BODY, size_pt=10, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Amount in words (full merge) ──────────────────────────────────────────
    r_words = tbl.rows[R_WORDS]
    _rw = list(r_words.cells)
    cw  = _rw[0].merge(_rw[9])
    _cell_width(cw, sum(COL_W))
    _cell_valign(cw, "center")
    words = _num_to_words(int(round(grand_total)))
    _para(cw,
          f"Order Value in Words:  {words} Rupees Only",
          font=F_BODY, size_pt=10, bold=True, italic=True,
          align=WD_ALIGN_PARAGRAPH.LEFT)

    return tbl


# ═════════════════════════════════════════════════════════════════════════════
# Amount in words helper
# ═════════════════════════════════════════════════════════════════════════════

def _num_to_words(n):
    ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight",
            "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen",
            "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty",
            "Sixty", "Seventy", "Eighty", "Ninety"]
    def _h(n):
        if n == 0:  return ""
        if n < 20:  return ones[n] + " "
        if n < 100: return tens[n // 10] + (" " + ones[n % 10] if n % 10 else "") + " "
        return ones[n // 100] + " Hundred " + _h(n % 100)
    n = int(n)
    if n == 0: return "Zero"
    crore = n // 10000000; n %= 10000000
    lakh  = n // 100000;   n %= 100000
    thou  = n // 1000;     n %= 1000
    parts = []
    if crore: parts.append(_h(crore).strip() + " Crore")
    if lakh:  parts.append(_h(lakh).strip()  + " Lakh")
    if thou:  parts.append(_h(thou).strip()  + " Thousand")
    if n:     parts.append(_h(n).strip())
    return " ".join(parts)


# ═════════════════════════════════════════════════════════════════════════════
# Terms & Conditions section
# ═════════════════════════════════════════════════════════════════════════════

def _add_terms(doc, terms):
    if not terms:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("SPECIFIC TERMS & CONDITIONS")
    r.font.name = F_BODY; r.font.size = Pt(10); r.font.bold = True

    for i, t in enumerate(terms, 1):
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(2)
        tp.paragraph_format.left_indent  = Pt(14)
        tp.paragraph_format.first_line_indent = Pt(-14)
        hr = tp.add_run(f"{i}. {t.get('header', '')}: ")
        hr.font.name = F_BODY; hr.font.size = Pt(10); hr.font.bold = True
        dr = tp.add_run(t.get("description", ""))
        dr.font.name = F_BODY; dr.font.size = Pt(10); dr.font.bold = False
        tp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ═════════════════════════════════════════════════════════════════════════════
# Signature block
# ═════════════════════════════════════════════════════════════════════════════

def _add_signature(doc, prepared_by, checked_by, authorised_by):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("For DISHAAN HI-TECH (INDIA) PVT. LTD.")
    r.font.name = F_BODY; r.font.size = Pt(10); r.font.bold = True

    # 3-column signature table
    sig_tbl = doc.add_table(rows=2, cols=3)
    sig_tbl.autofit = False; sig_tbl.allow_autofit = False
    _set_table_borders(sig_tbl)
    col_w = sum(COL_W) // 3

    slots = [
        ("Prepared By",   prepared_by),
        ("Checked By",    checked_by[0] if checked_by else None),
        ("Authorised By", authorised_by),
    ]
    for ci, (label, person) in enumerate(slots):
        hdr_cell  = sig_tbl.rows[0].cells[ci]
        body_cell = sig_tbl.rows[1].cells[ci]
        _cell_width(hdr_cell,  col_w)
        _cell_width(body_cell, col_w)
        _cell_shading(hdr_cell, C_GRAY_HD)
        _cell_valign(hdr_cell,  "center")
        _cell_valign(body_cell, "center")
        _para(hdr_cell, label, font=F_BODY, size_pt=9, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=3, sp_after=3)
        name_val = f"{person['name']}\n{person['at']}" if person else "—"
        _para(body_cell, name_val, font=F_BODY, size_pt=9,
              align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=15, sp_after=15)


# ═════════════════════════════════════════════════════════════════════════════
# Full document builder
# ═════════════════════════════════════════════════════════════════════════════

def _build_doc(order, company, vendor, items, terms,
               prepared_by, checked_by, authorised_by,
               qr_bytes, fp, generated_at):
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Twips(_PG_W)
    sec.page_height   = Twips(_PG_H)
    sec.top_margin    = Twips(_MG_T)
    sec.bottom_margin = Twips(_MG_B)
    sec.left_margin   = Twips(_MG_L)
    sec.right_margin  = Twips(_MG_R)

    # Default paragraph spacing = 0
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].font.name = F_BODY
    doc.styles["Normal"].font.size = Pt(9)

    _build_word_header(doc, qr_bytes, company)
    _build_word_footer(doc, generated_at, fp)

    # ── GSTIN line ────────────────────────────────────────────────────────────
    gp = doc.add_paragraph()
    gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gp.paragraph_format.space_before = Pt(2)
    gp.paragraph_format.space_after  = Pt(4)
    gr = gp.add_run(f"GSTIN NO.: {company.gstn or '—'}")
    gr.font.name = F_GSTIN; gr.font.size = Pt(10)

    # ── Main data table ───────────────────────────────────────────────────────
    _build_main_table(doc, order, vendor, company, items)

    # ── Terms & Conditions ────────────────────────────────────────────────────
    _add_terms(doc, terms)

    # ── Signature ─────────────────────────────────────────────────────────────
    _add_signature(doc, prepared_by, checked_by, authorised_by)

    return doc


# ═════════════════════════════════════════════════════════════════════════════
# Public API
# ═════════════════════════════════════════════════════════════════════════════

def generate_order_pdf(order_id, base_url, force=False):
    try:
        order = OrderMaster.query.get(order_id)
        if not order:
            return res("Order not found", [], 404)

        if force and getattr(order, "pdf_url", None):
            # Delete existing PDF file from disk before regenerating
            try:
                base_prefix = current_app.config.get("PDF_BASE_URL", "/resource/order/pdf-file")
                relative    = order.pdf_url.replace(base_prefix + "/", "", 1)
                old_path    = os.path.join(_storage_root(), relative)
                if os.path.exists(old_path):
                    os.remove(old_path)
            except Exception:
                pass
            order.pdf_url = order.pdf_token = order.pdf_generated_at = None

        if (not force
                and getattr(order, "pdf_url",          None)
                and getattr(order, "pdf_generated_at", None)):
            age = (datetime.utcnow() - order.pdf_generated_at).days
            if age < PDF_EXPIRY_DAYS:
                base_prefix = current_app.config.get("PDF_BASE_URL", "/resource/order/pdf-file")
                relative = order.pdf_url.replace(base_prefix + "/", "", 1)
                return serve_pdf_file(relative)

        company = Company.query.first()
        fp      = _fingerprint(order)
        token   = _serializer().dumps(
            {"oid": order.id, "ono": order.order_no, "fp": fp},
            salt=PDF_TOKEN_SALT,
        )
        verify_url   = f"{base_url.rstrip('/')}/resource/order/verify/{token}"
        qr_bytes     = _qr_with_brackets(verify_url)
        generated_at = datetime.utcnow().strftime("%d-%b-%Y %H:%M UTC")

        # Vendor
        vendor = order.vendor
        if not vendor:
            class _FV:
                ledger_name = registered_address = pan = gstin = ""
                state_name = state_code = primary_contact_person = ""
                primary_contact_number = email = ""
            vendor = _FV()

        # Items
        items = []
        for i, oi in enumerate(order.items, 1):
            hsn = getattr(oi.item, "hsn_sac", "") or "" if oi.item else ""
            items.append({
                "sl":          i,
                "item_code":   oi.item_code or "",
                "item_name":   oi.item.item_name if oi.item else "",
                "note":        oi.custom_note or "",
                "hsn_sac":     hsn,
                "unit":        oi.item.unit.unit_name if oi.item and oi.item.unit else "",
                "qty":         float(oi.qty or 0),
                "rate":        float(oi.rate or 0),
                "amount":      float(oi.amount or 0),
                "gst_percent": float(oi.gst_percent or 0),
                "gst_amount":  float(oi.gst_amount or 0),
                "total":       float(oi.amount or 0) + float(oi.gst_amount or 0),
            })

        # Terms
        terms = [
            {
                "header":      t.term.header if t.term else "",
                "description": t.custom_description or (t.term.term_description if t.term else ""),
            }
            for t in order.terms_conditions
        ]

        # Approval history
        history = (
            db.session.query(ApprovalHistory, User.username)
            .outerjoin(User, User.id == ApprovalHistory.action_by)
            .filter(ApprovalHistory.module_code == "order",
                    ApprovalHistory.record_id   == order.id)
            .order_by(ApprovalHistory.id.asc())
            .all()
        )
        prepared_by = None; checked_by = []; authorised_by = None
        for h, username in history:
            name = username or "—"
            at   = h.created_at.strftime("%d-%b-%Y; %H:%M") if h.created_at else ""
            if   h.action == "SUBMIT":        prepared_by   = {"name": name, "at": at}
            elif h.action == "FINAL_APPROVE": authorised_by = {"name": name, "at": at}
            elif h.action == "APPROVE":       checked_by.append({"name": name, "at": at})

        doc = _build_doc(order, company, vendor, items, terms,
                         prepared_by, checked_by, authorised_by,
                         qr_bytes, fp, generated_at)

        # DOCX → PDF: LibreOffice first, fallback to docx2pdf on Windows
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "order.docx")
            doc.save(docx_path)
            # Debug: copy DOCX to Downloads for inspection
            import shutil
            _dbg = r"C:\Users\SoumyajitMaity\Downloads\debug_order.docx"
            shutil.copy2(docx_path, _dbg)
            pdf_path  = os.path.join(tmpdir, "order.pdf")
            try:
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf",
                     "--outdir", tmpdir, docx_path],
                    check=True, timeout=60,
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            except (FileNotFoundError, subprocess.CalledProcessError):
                import pythoncom
                from docx2pdf import convert
                pythoncom.CoInitialize()
                try:
                    convert(docx_path, pdf_path)
                finally:
                    pythoncom.CoUninitialize()
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

        pdf_url = _save_pdf(pdf_bytes, order.order_no)
        order.pdf_url          = pdf_url
        order.pdf_token        = token
        order.pdf_generated_at = datetime.utcnow()
        db.session.commit()

        response = make_response(pdf_bytes)
        response.headers["Content-Type"]        = "application/pdf"
        response.headers["Content-Disposition"] = f'inline; filename="{order.order_no}.pdf"'
        return response

    except Exception as e:
        db.session.rollback()
        return res(str(e), [], 500)


def verify_order_pdf(token):
    try:
        data  = _serializer().loads(token, salt=PDF_TOKEN_SALT,
                                    max_age=PDF_EXPIRY_DAYS * 86400)
        order = OrderMaster.query.get(data["oid"])
        if not order or order.pdf_token != token:
            return _page("invalid")
        if data.get("fp") != _fingerprint(order):
            return _page("tampered", order)
        base_prefix = current_app.config.get("PDF_BASE_URL", "/resource/order/pdf-file")
        relative = order.pdf_url.replace(base_prefix + "/", "", 1)
        return serve_pdf_file(relative)
    except SignatureExpired as e:
        try:
            payload = e.payload
            if payload:
                order = OrderMaster.query.get(payload.get("oid"))
                if order and order.pdf_token == token:
                    order.pdf_url = order.pdf_token = order.pdf_generated_at = None
                    db.session.commit()
        except Exception:
            pass
        return _page("expired")
    except BadSignature:
        return _page("invalid")
    except Exception as e:
        return _page("error", msg=str(e))


def _page(status, order=None, msg=""):
    cfg = {
        "valid":    ("#2e7d32", "✓ DOCUMENT VERIFIED",
                     f"Order <b>{order.order_no}</b> is authentic.<br>"
                     f"Total: ₹{float(order.total_amount or 0):,.2f}" if order else "Verified."),
        "tampered": ("#b71c1c", "⚠ TAMPERED DOCUMENT",
                     "Data does not match the original.<br><b>Do not accept this document.</b>"),
        "expired":  ("#e65100", "⏱ QR CODE EXPIRED",
                     "This QR is valid for 2 days only.<br>Request a regenerated document."),
        "invalid":  ("#37474f", "✗ INVALID QR", "This QR code is not recognised."),
        "error":    ("#37474f", "Error", msg),
    }
    color, title, body = cfg.get(status, cfg["invalid"])
    code = 200 if status == "valid" else 410 if status == "expired" else 400
    return make_response(f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>Order Verification</title>
<style>
  body{{font-family:Arial,sans-serif;display:flex;justify-content:center;
       align-items:center;min-height:100vh;margin:0;background:#f5f5f5;}}
  .card{{background:#fff;border-radius:12px;padding:40px;max-width:480px;
         text-align:center;box-shadow:0 4px 20px rgba(0,0,0,.12);}}
  .icon{{font-size:52px;color:{color};}}
  .title{{font-size:20px;font-weight:bold;color:{color};margin:12px 0;}}
  .body{{font-size:14px;color:#555;line-height:1.7;}}
</style></head><body>
<div class="card">
  <div class="icon">{title[0]}</div>
  <div class="title">{title}</div>
  <div class="body">{body}</div>
</div></body></html>""", code)
